"""
Stop hook — runs when Claude Code finishes responding.

Does the mechanical / deterministic parts of the Session Protocol:
  1. Backs up CLAUDE.md to CLAUDE.md.backup
  2. Finds the newest .docx in the devlog directory
  3. Captures `git log --oneline -10`
  4. Rewrites the ## Last Devlog section with the new snapshot
  5. Refreshes ## Infrastructure State lines whose values appear in the
     devlog's "Key Infrastructure References" table
  6. Writes .claude/pending_session_summary.json so the NEXT session's
     SessionStart hook can tell the model to synthesize ## Key Learnings
     from the just-ended transcript (we can't do model work here — this
     script is a shell command, not a Claude session).

Must never fail the session. All errors are swallowed and logged to
.claude/hooks/stop_hook.log — the hook always exits 0.

Config via env:
  PSITTA_DEVLOG_DIR  — override the devlog directory (default is the
                       Windows OneDrive path from CLAUDE.md).
"""
from __future__ import annotations

import json
import os
import re
import subprocess
import sys
import traceback
from datetime import datetime, timezone
from pathlib import Path

HOOK_DIR = Path(__file__).resolve().parent
CLAUDE_DIR = HOOK_DIR.parent
REPO_ROOT = CLAUDE_DIR.parent
CLAUDE_MD = REPO_ROOT / "CLAUDE.md"
BACKUP_MD = REPO_ROOT / "CLAUDE.md.backup"
PENDING_MARKER = CLAUDE_DIR / "pending_session_summary.json"
LOG_FILE = HOOK_DIR / "stop_hook.log"

DEFAULT_DEVLOG_DIR = Path(
    os.environ.get(
        "PSITTA_DEVLOG_DIR",
        r"C:\Users\Admin\OneDrive\_Psitta\Docs\DevLogs",
    )
)


def log(msg: str) -> None:
    ts = datetime.now(timezone.utc).isoformat(timespec="seconds")
    try:
        with LOG_FILE.open("a", encoding="utf-8") as f:
            f.write(f"[{ts}] {msg}\n")
    except Exception:
        pass


def read_hook_event() -> dict:
    try:
        raw = sys.stdin.read()
        if not raw.strip():
            return {}
        return json.loads(raw)
    except Exception as e:
        log(f"failed to parse stdin event: {e}")
        return {}


def newest_devlog(devlog_dir: Path) -> Path | None:
    if not devlog_dir.exists():
        log(f"devlog dir missing: {devlog_dir}")
        return None
    docs = sorted(
        devlog_dir.glob("*.docx"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    return docs[0] if docs else None


def extract_devlog(docx_path: Path) -> dict:
    """Return {title, date, focus, infra: {item: value}} from the devlog."""
    try:
        from docx import Document  # type: ignore
    except ImportError:
        log("python-docx not installed — skipping devlog extraction")
        return {}
    try:
        d = Document(str(docx_path))
    except Exception as e:
        log(f"could not open devlog {docx_path.name}: {e}")
        return {}

    out: dict = {"title": None, "date": None, "focus": None, "infra": {}}
    for p in d.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if out["title"] is None and "Development Log" in text:
            out["title"] = text
        if out["date"] is None:
            m = re.search(
                r"(January|February|March|April|May|June|July|August|"
                r"September|October|November|December)\s+\d{1,2},\s+\d{4}",
                text,
            )
            if m:
                out["date"] = m.group(0)

    for table in d.tables:
        headers = [c.text.strip().lower() for c in table.rows[0].cells]
        # Key Infrastructure References table: first header is item/value
        if headers and headers[0] in {"item", "field"} and len(headers) >= 2:
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2 and cells[0]:
                    key = cells[0]
                    value = cells[1]
                    if key.lower() == "session focus":
                        out["focus"] = value
                    else:
                        out["infra"][key] = value
    return out


def git_log_oneline(n: int = 10) -> str:
    try:
        env = os.environ.copy()
        env["LC_ALL"] = "C.UTF-8"
        env["PYTHONIOENCODING"] = "utf-8"
        res = subprocess.run(
            ["git", "log", "--oneline", f"-{n}"],
            cwd=str(REPO_ROOT),
            capture_output=True,
            timeout=10,
            env=env,
        )
        if res.returncode == 0:
            return res.stdout.decode("utf-8", errors="replace").strip()
        log(
            "git log non-zero: "
            + res.stderr.decode("utf-8", errors="replace").strip()
        )
    except Exception as e:
        log(f"git log failed: {e}")
    return ""


def replace_section(content: str, heading: str, new_body: str) -> str:
    """
    Replace the body of a `## heading` section with new_body. The section
    runs from the heading line up to (but not including) the next `## ` at
    column 0. If the heading isn't found, content is returned unchanged.
    """
    pattern = re.compile(
        rf"(^##\s+{re.escape(heading)}\s*\n)(.*?)(?=^##\s|\Z)",
        re.DOTALL | re.MULTILINE,
    )
    if not pattern.search(content):
        return content
    return pattern.sub(lambda m: m.group(1) + new_body, content)


def update_infrastructure_lines(content: str, infra: dict) -> str:
    """
    Conservative update: for each bullet line in ## Infrastructure State,
    if the bullet's bold label matches a key in infra, rewrite that line.
    Unknown lines are left alone. New keys are not added.
    """
    if not infra:
        return content

    # Map devlog keys to CLAUDE.md bullet labels (both sides hand-picked).
    key_map = {
        "Azure Speech Region": "Azure Speech Region",
        "Secrets Manager": "Secrets Manager",
        "AWS Account": "AWS Account",
        "Cognito User Pool": "Cognito User Pool",
        "Production API": "Production API",
        "ECS Cluster/Service": "ECS Cluster/Service",
        "GitHub Repo": "GitHub",
    }

    def strip_md(s: str) -> str:
        return s.replace("`", "").lower()

    def rewrite(line: str) -> str:
        m = re.match(r"^(-\s+\*\*)([^*]+)(\*\*:\s*)(.*)$", line)
        if not m:
            return line
        label = m.group(2).strip()
        devlog_key = next(
            (k for k, v in key_map.items() if v == label), None
        )
        if not devlog_key or devlog_key not in infra:
            return line
        new_value = infra[devlog_key]
        current_value = m.group(4)
        # Heuristic: devlog values always lead with the primary identifier
        # (region name, secret path, pool ID, hostname, etc). If that
        # first token is already in the current line, the line is
        # already accurate — don't downgrade the human-written context.
        first_token = new_value.split()[0] if new_value.split() else ""
        if first_token and strip_md(first_token) in strip_md(current_value):
            return line
        return f"{m.group(1)}{m.group(2)}{m.group(3)}{new_value}"

    section_re = re.compile(
        r"(^##\s+Infrastructure State\s*\n)(.*?)(?=^##\s|\Z)",
        re.DOTALL | re.MULTILINE,
    )
    m = section_re.search(content)
    if not m:
        return content
    body = m.group(2)
    original_lines = body.splitlines(keepends=True)
    new_lines = [
        (rewrite(ln.rstrip("\r\n")) + ln[len(ln.rstrip("\r\n")):])
        for ln in original_lines
    ]
    new_body = "".join(new_lines)
    if new_body == body:
        return content
    return content[: m.start(2)] + new_body + content[m.end(2):]


def build_last_devlog_body(
    devlog_path: Path | None,
    meta: dict,
    git_log: str,
) -> str:
    lines: list[str] = []
    if devlog_path is not None:
        lines.append(f"- **File**: `{devlog_path}`")
    if meta.get("date"):
        lines.append(f"- **Date**: {meta['date']}")
    if meta.get("title"):
        lines.append(f"- **Title**: {meta['title']}")
    if meta.get("focus"):
        lines.append(f"- **Focus**: {meta['focus']}")
    if git_log:
        lines.append("- **Recent commits** (`git log --oneline -10`):")
        lines.append("")
        lines.append("```")
        lines.append(git_log)
        lines.append("```")
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    lines.append(f"- _Auto-updated by Stop hook at {ts}_")
    lines.append("")
    return "\n".join(lines) + "\n"


def write_pending_marker(event: dict, devlog_path: Path | None, git_log: str) -> None:
    payload = {
        "written_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "session_id": event.get("session_id"),
        "transcript_path": event.get("transcript_path"),
        "newest_devlog": str(devlog_path) if devlog_path else None,
        "git_log_oneline_10": git_log,
        "repo_root": str(REPO_ROOT),
    }
    # Preserve prior markers if one already exists (rare — previous session
    # never had its SessionStart follow-up). Merge into a list.
    if PENDING_MARKER.exists():
        try:
            prior = json.loads(PENDING_MARKER.read_text(encoding="utf-8"))
            if isinstance(prior, dict):
                prior = [prior]
            prior.append(payload)
            PENDING_MARKER.write_text(
                json.dumps(prior, indent=2), encoding="utf-8"
            )
            return
        except Exception as e:
            log(f"could not merge prior marker, overwriting: {e}")
    PENDING_MARKER.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def main() -> int:
    try:
        event = read_hook_event()
        # Avoid infinite loops if Claude Code re-invokes the stop hook.
        if event.get("stop_hook_active"):
            log("stop_hook_active=true — bailing out")
            return 0

        if not CLAUDE_MD.exists():
            log(f"CLAUDE.md missing at {CLAUDE_MD} — nothing to update")
            return 0

        original = CLAUDE_MD.read_text(encoding="utf-8")
        BACKUP_MD.write_text(original, encoding="utf-8")

        devlog_path = newest_devlog(DEFAULT_DEVLOG_DIR)
        meta = extract_devlog(devlog_path) if devlog_path else {}
        git_log = git_log_oneline(10)

        updated = original
        new_body = build_last_devlog_body(devlog_path, meta, git_log)
        updated = replace_section(updated, "Last Devlog", new_body)
        updated = update_infrastructure_lines(updated, meta.get("infra", {}))

        if updated != original:
            CLAUDE_MD.write_text(updated, encoding="utf-8")
            log(
                f"CLAUDE.md updated (devlog="
                f"{devlog_path.name if devlog_path else 'none'}, "
                f"git_log_lines={len(git_log.splitlines())})"
            )
        else:
            log("no changes to CLAUDE.md")

        write_pending_marker(event, devlog_path, git_log)
        log("pending marker written")
        return 0
    except Exception as e:
        log(f"unhandled error: {e}\n{traceback.format_exc()}")
        return 0


if __name__ == "__main__":
    sys.exit(main())
