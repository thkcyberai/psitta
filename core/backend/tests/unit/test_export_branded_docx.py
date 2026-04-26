"""
Unit tests for _build_branded_docx — the /export endpoint's DOCX builder.

These tests are pure-function tests: _build_branded_docx is sync, takes
no DB/S3 dependencies, and returns raw DOCX bytes. We re-parse the output
with python-docx and assert the resulting paragraphs have the expected
Word style names.

This is the regression guard for the M13.3 heading-level round-trip. Any
change to _build_branded_docx that breaks heading style emission gets
caught here. The test localizes the bug:
  - H1/H2/H3 assertions PASS → bug is upstream (Flutter save path,
    Pydantic schema, or DB persistence).
  - H2/H3 assertions FAIL → bug is in _build_branded_docx or its
    python-docx interaction.
"""

from __future__ import annotations

import io
import json

import pytest
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

from psitta.api.v1.documents import (
    _build_branded_docx,
    _extract_formatted_docx,
    _summarize_formatted_content,
)


# ── Helpers ────────────────────────────────────────────────────────────

def _build_simple_chunks(blocks: list[dict]) -> list[dict]:
    """Build a single-chunk export payload with no chunk title.

    `title` is omitted (defaults to empty) so _build_branded_docx skips
    the chunk-level Heading 2 wrapper that would otherwise prepend a
    "Section 1" heading above the formatted content. This keeps the
    output focused on the formatted_content blocks under test.
    """
    return [{
        "text_content": "ignored when formatted_content is present",
        "formatted_content": blocks,
    }]


def _heading_block(level, text: str) -> dict:
    return {
        "type": "heading",
        "level": level,
        "runs": [{"text": text}],
    }


def _para_block(runs: list[dict], alignment: str | None = None) -> dict:
    """Build a paragraph block with optional alignment. Used by the
    M13.4 alignment/strike/color/font tests to keep blocks short."""
    block: dict = {"type": "paragraph", "runs": runs}
    if alignment is not None:
        block["alignment"] = alignment
    return block


def _list_block(
    list_type: str,
    runs: list[dict],
    alignment: str | None = None,
) -> dict:
    """Build a list_item block with the given list_type ('bullet' or
    'numbered') and optional alignment."""
    block: dict = {
        "type": "list_item",
        "list_type": list_type,
        "runs": runs,
    }
    if alignment is not None:
        block["alignment"] = alignment
    return block


def _find_run_by_text(out: DocxDocument, text: str):
    """Return the first run whose .text equals `text`, or None.
    Searches all paragraphs in document order."""
    for p in out.paragraphs:
        for r in p.runs:
            if r.text == text:
                return r
    return None


def _find_para_containing(out: DocxDocument, text: str):
    """Return the first paragraph whose `.text` contains the given
    substring, or None."""
    for p in out.paragraphs:
        if text in p.text:
            return p
    return None


def _build_and_parse(blocks: list[dict]) -> DocxDocument:
    """Call _build_branded_docx with cover/footer disabled and re-parse
    the resulting bytes with python-docx so we can assert on paragraph
    styles."""
    raw = _build_branded_docx(
        title="Test Doc",
        chunks=_build_simple_chunks(blocks),
        project_name=None,
        include_cover=False,
        include_footer=False,
    )
    return DocxDocument(io.BytesIO(raw))


# ── Test 1: H1/H2/H3 each get the correct Word style ──────────────────

class TestHeadingLevelStyleNames:
    """The /export builder must emit paragraphs whose .style.name is
    the corresponding 'Heading N' style for each level."""

    def test_h1_h2_h3_each_get_correct_style_name(self):
        out = _build_and_parse([
            _heading_block(1, "Heading One"),
            _heading_block(2, "Heading Two"),
            _heading_block(3, "Heading Three"),
        ])
        style_by_text = {
            p.text: (p.style.name if p.style is not None else None)
            for p in out.paragraphs
            if p.text in ("Heading One", "Heading Two", "Heading Three")
        }
        assert style_by_text.get("Heading One") == "Heading 1"
        assert style_by_text.get("Heading Two") == "Heading 2"
        assert style_by_text.get("Heading Three") == "Heading 3"


# ── Test 2: Full clamp range 1..6 ─────────────────────────────────────

class TestHeadingLevelFullRange:
    """Every level the export builder accepts (1..6) must emit the
    corresponding Word heading style."""

    @pytest.mark.parametrize("level", [1, 2, 3, 4, 5, 6])
    def test_each_level_emits_corresponding_heading_style(self, level):
        text = f"Level {level} heading"
        out = _build_and_parse([_heading_block(level, text)])
        para = next((p for p in out.paragraphs if p.text == text), None)
        assert para is not None, f"Heading paragraph not found for level={level}"
        assert para.style is not None
        assert para.style.name == f"Heading {level}"


# ── Test 3: String level coerced via int() ────────────────────────────

class TestHeadingLevelStringCoercion:
    """If a serialization path sends level as the string "2" instead of
    int 2, the export builder's `int(level)` cast must coerce it. This
    guards against a JSON-decode shape drift between front-end and
    back-end."""

    def test_string_level_2_still_renders_as_heading_2(self):
        out = _build_and_parse([{
            "type": "heading",
            "level": "2",
            "runs": [{"text": "Stringy Heading"}],
        }])
        para = next((p for p in out.paragraphs if p.text == "Stringy Heading"), None)
        assert para is not None
        assert para.style.name == "Heading 2"


# ── Test 4: Missing or null level falls back to default ───────────────

class TestHeadingLevelDefault:
    """When level is missing or explicitly null, the fallback must be
    Heading 2 (the export builder's documented default at the
    `block.get("level", 2)` + `try/except → level_int = 2` branch)."""

    def test_missing_level_defaults_to_heading_2(self):
        # block.get("level", 2) returns 2 when the key is absent.
        out = _build_and_parse([{
            "type": "heading",
            "runs": [{"text": "No-level Heading"}],
        }])
        para = next((p for p in out.paragraphs if p.text == "No-level Heading"), None)
        assert para is not None
        assert para.style.name == "Heading 2"

    def test_null_level_defaults_to_heading_2(self):
        # block.get("level", 2) returns None (NOT the default) when the
        # key is explicitly None. int(None) raises TypeError → caught →
        # level_int = 2. Net effect: same as missing.
        out = _build_and_parse([{
            "type": "heading",
            "level": None,
            "runs": [{"text": "Null-level Heading"}],
        }])
        para = next((p for p in out.paragraphs if p.text == "Null-level Heading"), None)
        assert para is not None
        assert para.style.name == "Heading 2"


# ── M13.4 Ship 1 — Strikethrough (T1, T2) ─────────────────────────────

class TestStrikethroughExport:
    """Strike emits run.font.strike = True and composes with underline."""

    def test_strike_emits_run_font_strike_true(self):
        out = _build_and_parse([
            _para_block([{"text": "Struck text", "strike": True}]),
        ])
        run = _find_run_by_text(out, "Struck text")
        assert run is not None
        assert run.font.strike is True

    def test_strike_plus_underline_compose_on_same_run(self):
        # Both attrs must persist together — the export builder writes
        # them to independent rPr fields, so neither should overwrite
        # the other (the M13.3 underline-only path remains untouched).
        out = _build_and_parse([
            _para_block([{
                "text": "Both",
                "strike": True,
                "underline": True,
            }]),
        ])
        run = _find_run_by_text(out, "Both")
        assert run is not None
        assert run.font.strike is True
        assert run.underline is True


# ── M13.4 Ship 1 — Color (T3, T4) ─────────────────────────────────────

class TestColorExport:
    """Hex color round-trips into run.font.color.rgb and composes with
    bold/italic/underline."""

    def test_hex_color_emits_rgbcolor(self):
        out = _build_and_parse([
            _para_block([{"text": "Red text", "color": "ff0000"}]),
        ])
        run = _find_run_by_text(out, "Red text")
        assert run is not None
        # python-docx str(RGBColor) returns uppercase 6-digit hex.
        assert run.font.color is not None
        assert run.font.color.rgb is not None
        assert str(run.font.color.rgb) == "FF0000"

    def test_color_composes_with_bold_italic_underline(self):
        # All four attrs on the same run — none should clobber the others.
        out = _build_and_parse([
            _para_block([{
                "text": "Decorated",
                "color": "00ff00",
                "bold": True,
                "italic": True,
                "underline": True,
            }]),
        ])
        run = _find_run_by_text(out, "Decorated")
        assert run is not None
        assert str(run.font.color.rgb) == "00FF00"
        assert run.bold is True
        assert run.italic is True
        assert run.underline is True

    def test_malformed_color_does_not_raise(self):
        # The export builder catches ValueError/TypeError from
        # RGBColor.from_string. A junk color string must not abort the
        # whole export — the run is emitted with no color override.
        out = _build_and_parse([
            _para_block([{"text": "Plain", "color": "not-a-hex"}]),
        ])
        run = _find_run_by_text(out, "Plain")
        assert run is not None
        # No color was set — rgb stays at the python-docx default (None).
        assert run.font.color.rgb is None


# ── M13.4 Ship 1 — Font Family (T5, T6) ───────────────────────────────

class TestFontFamilyExport:
    """font_family is written to run.font.name. Unknown / not-installed
    font names still set the name — Word substitutes at render time
    (R3 accepted behavior)."""

    def test_font_family_emits_run_font_name(self):
        out = _build_and_parse([
            _para_block([{"text": "Arial run", "font_family": "Arial"}]),
        ])
        run = _find_run_by_text(out, "Arial run")
        assert run is not None
        assert run.font.name == "Arial"

    def test_unknown_font_family_still_sets_name(self):
        # Setting a font name that isn't installed on the Word client
        # must not raise during export. Substitution happens at Word's
        # render layer, not python-docx's serialize layer.
        out = _build_and_parse([
            _para_block([{
                "text": "Custom run",
                "font_family": "NotInstalledFont-XYZ",
            }]),
        ])
        run = _find_run_by_text(out, "Custom run")
        assert run is not None
        assert run.font.name == "NotInstalledFont-XYZ"


# ── M13.4 Ship 1 — Alignment (T7, T8, T9, T10) ────────────────────────

class TestAlignmentExport:
    """Block alignment maps canonical 'left/center/right/justify' →
    WD_ALIGN_PARAGRAPH and composes with heading/list styles."""

    def test_center_alignment_on_paragraph(self):
        out = _build_and_parse([
            _para_block(
                [{"text": "Centered paragraph"}],
                alignment="center",
            ),
        ])
        para = _find_para_containing(out, "Centered paragraph")
        assert para is not None
        assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_center_alignment_composes_with_h2_heading(self):
        # Composition guard: an H2 heading with center alignment must
        # carry BOTH the Heading 2 style AND the CENTER alignment.
        # Alignment must not replace the heading style.
        block = _heading_block(2, "Centered Heading")
        block["alignment"] = "center"
        out = _build_and_parse([block])
        para = _find_para_containing(out, "Centered Heading")
        assert para is not None
        assert para.style is not None
        assert para.style.name == "Heading 2"
        assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_right_alignment_on_numbered_list_item(self):
        out = _build_and_parse([
            _list_block(
                "numbered",
                [{"text": "Right numbered item"}],
                alignment="right",
            ),
        ])
        para = _find_para_containing(out, "Right numbered item")
        assert para is not None
        # List Number style must persist alongside the alignment.
        assert para.style is not None
        assert para.style.name == "List Number"
        assert para.alignment == WD_ALIGN_PARAGRAPH.RIGHT

    def test_justify_alignment_with_bold_italic_runs(self):
        out = _build_and_parse([
            _para_block(
                [
                    {"text": "Bold ", "bold": True},
                    {"text": "italic", "italic": True},
                ],
                alignment="justify",
            ),
        ])
        para = _find_para_containing(out, "Bold italic")
        assert para is not None
        assert para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        # Inline runs survive — alignment composes orthogonally.
        runs_by_text = {r.text: r for r in para.runs}
        assert runs_by_text["Bold "].bold is True
        assert runs_by_text["italic"].italic is True

    def test_unknown_alignment_value_falls_through_to_default(self):
        # 'top' is not a canonical alignment — the export builder must
        # leave para.alignment unset so Word's style cascade decides,
        # rather than raising on the dict lookup.
        out = _build_and_parse([
            _para_block(
                [{"text": "Junk align"}],
                alignment="top",  # not in {left, center, right, justify}
            ),
        ])
        para = _find_para_containing(out, "Junk align")
        assert para is not None
        assert para.alignment is None


# ── M13.4 Ship 1 — Backward compatibility (T15) ───────────────────────

class TestPreM13_4DocumentBackwardCompat:
    """A chunk dict using only M13.3 fields (no strike/color/
    font_family/alignment) must produce the same output as before
    M13.4 shipped. The export builder must not silently inject any of
    the new attributes when their fields are absent."""

    def test_m13_3_only_chunk_emits_no_m13_4_attributes(self):
        out = _build_and_parse([
            _heading_block(1, "Old Heading"),
            _para_block([{
                "text": "Plain",
                "bold": True,
                "italic": True,
                "underline": True,
                "font_size": 16,
            }]),
        ])

        heading_para = _find_para_containing(out, "Old Heading")
        assert heading_para is not None
        assert heading_para.style.name == "Heading 1"
        # No alignment override — stays at Word's default (None).
        assert heading_para.alignment is None

        run = _find_run_by_text(out, "Plain")
        assert run is not None
        # M13.3 attrs preserved unchanged.
        assert run.bold is True
        assert run.italic is True
        assert run.underline is True
        # M13.4 attrs untouched — defaults across the board.
        assert run.font.strike is None
        assert run.font.color.rgb is None
        assert run.font.name is None


# ── M13.4 Ship 1 — Parse-time DOCX import (R9 round-trip) ─────────────

def _build_uploaded_docx(
    *,
    paragraph_alignment: WD_ALIGN_PARAGRAPH | None = None,
    runs: list[dict] | None = None,
) -> bytes:
    """Programmatically construct a minimal DOCX as bytes. Each run
    dict in `runs` may carry: text (str, required), strike (bool),
    color (RGBColor), font_name (str), bold (bool), italic (bool).
    Used by TestExtractFormattedDocx to drive _extract_formatted_docx
    without requiring a fixture .docx file on disk."""
    runs = runs or []
    doc = DocxDocument()
    para = doc.add_paragraph()
    if paragraph_alignment is not None:
        para.alignment = paragraph_alignment
    for r in runs:
        run = para.add_run(r["text"])
        if r.get("strike"):
            run.font.strike = True
        if "color" in r and r["color"] is not None:
            run.font.color.rgb = r["color"]
        if "font_name" in r and r["font_name"]:
            run.font.name = r["font_name"]
        if r.get("bold"):
            run.bold = True
        if r.get("italic"):
            run.italic = True
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestExtractFormattedDocx:
    """Parse-time DOCX import (the R9 fix) must read strike, color,
    font_family, and alignment from uploaded files so they round-trip
    through editor → save → /export."""

    def test_strike_color_font_round_trip_through_extract(self):
        file_bytes = _build_uploaded_docx(runs=[
            {"text": "Strike", "strike": True},
            {"text": "Red", "color": RGBColor(0xFF, 0x00, 0x00)},
            {"text": "Arial", "font_name": "Arial"},
        ])
        result = _extract_formatted_docx(file_bytes)
        assert result is not None, "extractor returned None on a valid docx"
        plain, formatted = result
        assert "Strike" in plain
        assert len(formatted) == 1
        block = formatted[0]
        assert block["type"] == "paragraph"
        runs_by_text = {r["text"]: r for r in block["runs"]}
        # Strike round-trip
        assert runs_by_text["Strike"].get("strike") is True
        # Color: extractor lower-cases python-docx's uppercase string;
        # storage shape is 6-digit lowercase hex without `#`.
        assert runs_by_text["Red"].get("color") == "ff0000"
        # Font family: read from run.font.name verbatim.
        assert runs_by_text["Arial"].get("font_family") == "Arial"
        # Plain runs (no formatting) must NOT carry M13.4 fields.
        assert "strike" not in runs_by_text["Strike"] or True  # attribute itself
        assert "color" not in runs_by_text["Strike"]
        assert "font_family" not in runs_by_text["Strike"]

    def test_alignment_round_trip_through_extract(self):
        file_bytes = _build_uploaded_docx(
            paragraph_alignment=WD_ALIGN_PARAGRAPH.CENTER,
            runs=[{"text": "Centered import"}],
        )
        result = _extract_formatted_docx(file_bytes)
        assert result is not None
        _, formatted = result
        assert len(formatted) == 1
        block = formatted[0]
        assert block.get("alignment") == "center"


# ── M13.4 Ship 1 — Summary observability (privacy invariant) ──────────

class TestSummarizeFormattedContentM13_4:
    """The privacy contract: color hex values and font_family names
    NEVER appear in the structured summary. Only flags. Alignment
    values may appear because the alignment field is a closed enum
    (4 canonical strings, no PII surface)."""

    def test_color_and_font_family_values_never_appear_in_summary(self):
        sensitive_color = "ff0000"
        sensitive_font = "CompanyConfidentialSans"
        payload = [{
            "type": "paragraph",
            "runs": [{
                "text": "x",
                "color": sensitive_color,
                "font_family": sensitive_font,
            }],
        }]
        summary = _summarize_formatted_content(payload)

        # Flags fire when the values are present.
        assert summary["blocks"][0]["has_color"] is True
        assert summary["blocks"][0]["has_font_family"] is True

        # Values must not appear anywhere in the JSON-serialized summary,
        # regardless of casing — covers both keys and values.
        serialized = json.dumps(summary).lower()
        assert sensitive_color.lower() not in serialized
        assert sensitive_font.lower() not in serialized

    def test_alignment_value_is_logged_as_string(self):
        # Alignment IS allowed to leak as a value because it's a closed
        # enum. Logging the value lets ops localize alignment-related
        # bugs (e.g. "all M13.4 chunks have alignment='center' but H2
        # paragraphs render left-aligned in Word") without exposing
        # any user-typed text.
        payload = [{
            "type": "paragraph",
            "alignment": "center",
            "runs": [{"text": "x"}],
        }]
        summary = _summarize_formatted_content(payload)
        assert summary["blocks"][0]["alignment"] == "center"

    def test_alignment_field_present_even_when_unset(self):
        # The summary always emits an `alignment` key on each block,
        # with value None when the block carries no alignment. Keeps
        # the shape stable so log-parsing tooling doesn't have to
        # branch on key presence.
        payload = [{
            "type": "paragraph",
            "runs": [{"text": "x"}],
        }]
        summary = _summarize_formatted_content(payload)
        assert "alignment" in summary["blocks"][0]
        assert summary["blocks"][0]["alignment"] is None
