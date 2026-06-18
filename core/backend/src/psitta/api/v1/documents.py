"""Psitta - Document Management Routes."""

from __future__ import annotations

import io
import json
import re
from datetime import datetime, timezone
from typing import Annotated
from uuid import UUID, uuid4

import pysbd
import structlog
from pathlib import Path
from fastapi import APIRouter, BackgroundTasks, Depends, Form, HTTPException, Query, Request, UploadFile, status
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from pydantic import BaseModel, Field
from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession

from psitta.dependencies import get_current_user_id, get_db_session
from psitta.middleware.auth import TokenClaims
from psitta.schemas.api import ChunkCreateRequest, ChunkResponse, ChunkUpdateRequest, ResynthesizeResponse
from psitta.services import audit_service

logger = structlog.get_logger(__name__)


def _reflow_pdf_text(text: str) -> str:
    """Reflow PDF-extracted text.
    Keeps paragraph breaks (blank lines) but merges single newlines into spaces.
    """
    if not text:
        return ""
    t = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = t.split("\n")
    out: list[str] = []
    buf = ""
    for line in lines:
        ln = line.strip()
        if not ln:
            if buf.strip():
                out.append(buf.strip())
                buf = ""
            out.append("")
            continue
        if buf.endswith("-"):
            buf = buf[:-1] + ln
        else:
            buf = (buf + " " + ln).strip() if buf else ln
    if buf.strip():
        out.append(buf.strip())
    # collapse 3+ blank lines to max 2
    joined = "\n".join(out)
    joined = re.sub(r"\n{3,}", "\n\n", joined)
    return joined


def _sanitize_text_for_db(text: str) -> str:
    """Remove bytes/chars that Postgres TEXT cannot store.
    - strips NULL bytes (\x00)
    - strips other control chars except \n \r \t
    """
    if not text:
        return ""
    text = text.replace("\x00", "")
    text = "".join(ch for ch in text if (ch >= " " or ch in "\n\r\t"))
    return text


def _normalize_sentence_compare(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip().casefold()


def _normalize_pdf_page_marker(line: str) -> str:
    stripped = line.strip()
    stripped = re.sub(r"^[\s\[\](){}<>#*~\-.:|]+", "", stripped)
    stripped = re.sub(r"[\s\[\](){}<>#*~\-.:|]+$", "", stripped)
    return stripped


def _looks_like_pdf_page_number_line(line: str, page_number: int | None = None) -> bool:
    stripped = _normalize_pdf_page_marker(line)
    if not stripped:
        return False
    if page_number is not None:
        if stripped == str(page_number):
            return True
        if re.fullmatch(rf"page\s+{page_number}(?:\s+of\s+\d+)?", stripped, re.IGNORECASE):
            return True
        if re.fullmatch(rf"{page_number}\s*/\s*\d+", stripped):
            return True
    if re.fullmatch(r"page\s+[ivxlcdm]{1,8}", stripped, re.IGNORECASE):
        return True
    if re.fullmatch(r"[ivxlcdm]{2,8}", stripped, re.IGNORECASE):
        return True
    return bool(re.fullmatch(r"page\s+\d+(?:\s+of\s+\d+)?", stripped, re.IGNORECASE))


def _strip_pdf_page_number_prefix(text: str, page_number: int | None = None) -> str:
    if not text:
        return ""

    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = normalized.split("\n")
    filtered: list[str] = []
    removed_prefix = False

    for idx, line in enumerate(lines):
        stripped = line.strip()
        if idx < 2 and not removed_prefix and _looks_like_pdf_page_number_line(stripped, page_number):
            removed_prefix = True
            continue
        filtered.append(line)

    cleaned = "\n".join(filtered).lstrip()
    if page_number is not None:
        cleaned = re.sub(
            rf"^(?:page\s+)?{page_number}(?:\s+of\s+\d+)?[\s:.\-]+",
            "",
            cleaned,
            count=1,
            flags=re.IGNORECASE,
        )
    return cleaned.strip()


def _strip_pdf_page_number_suffix(text: str, page_number: int | None = None) -> str:
    if not text:
        return ""

    cleaned = text.strip()
    if page_number is not None:
        cleaned = re.sub(
            rf"[\s:.\-]+(?:page\s+)?{page_number}(?:\s+of\s+\d+)?$",
            "",
            cleaned,
            count=1,
            flags=re.IGNORECASE,
        )
        cleaned = re.sub(
            rf"[\s:.\-]+{page_number}$",
            "",
            cleaned,
            count=1,
        )
    cleaned = re.sub(
        r"[\s:.\-]+page\s+\d+(?:\s+of\s+\d+)?$",
        "",
        cleaned,
        count=1,
        flags=re.IGNORECASE,
    )
    return cleaned.strip()


def _strip_pdf_page_number_edge_tokens(text: str, page_number: int | None = None) -> str:
    if not text:
        return ""

    cleaned = text.strip()
    wrapper = r"[\s\[\](){}<>#*~:|.\-]*"
    start_patterns = [
        rf"^{wrapper}page\s+\d+(?:\s+of\s+\d+)?{wrapper}\s+",
        rf"^{wrapper}[ivxlcdm]{{2,8}}{wrapper}\s+",
    ]
    end_patterns = [
        rf"\s+{wrapper}page\s+\d+(?:\s+of\s+\d+)?{wrapper}$",
        rf"\s+{wrapper}[ivxlcdm]{{2,8}}{wrapper}$",
    ]
    if page_number is not None:
        start_patterns.insert(
            0,
            rf"^{wrapper}(?:page\s+)?{page_number}(?:\s+of\s+\d+)?{wrapper}\s+",
        )
        end_patterns.insert(
            0,
            rf"\s+{wrapper}(?:page\s+)?{page_number}(?:\s+of\s+\d+)?{wrapper}$",
        )

    for pattern in start_patterns:
        cleaned = re.sub(pattern, "", cleaned, count=1, flags=re.IGNORECASE)
    for pattern in end_patterns:
        cleaned = re.sub(pattern, "", cleaned, count=1, flags=re.IGNORECASE)
    return cleaned.strip()


def _strip_pdf_page_number_lines(text: str, page_number: int | None = None) -> str:
    if not text:
        return ""

    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = normalized.split("\n")
    filtered: list[str] = []
    last_index = len(lines) - 1

    for idx, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            filtered.append(line)
            continue
        near_edge = idx < 2 or idx >= max(0, last_index - 1)
        if near_edge and _looks_like_pdf_page_number_line(stripped, page_number):
            continue
        filtered.append(line)

    return "\n".join(filtered)


def _drop_duplicate_leading_sentence(text: str) -> str:
    if not text:
        return ""

    segmenter = pysbd.Segmenter(language="en", clean=False)
    sentences = [s for s in segmenter.segment(text) if s and s.strip()]
    if len(sentences) < 2:
        return text.strip()

    first = sentences[0].strip()
    second = sentences[1].strip()
    if _normalize_sentence_compare(first) != _normalize_sentence_compare(second):
        return text.strip()

    start = text.find(second)
    if start <= 0:
        return text.strip()
    return text[start:].lstrip()


def _add_terminal_punctuation(text: str) -> str:
    """Add a period to lines that look like headings/titles without terminal punctuation.
    Only applies to short lines (<=80 chars) to avoid adding periods to
    soft-wrapped body text lines from PDF extraction.
    """
    terminal = frozenset({".", "!", "?", ":", ";", ","})
    result = []
    for line in text.split("\n"):
        stripped = line.rstrip()
        # Only add period to short lines — long lines are wrapped body text
        if stripped and len(stripped) <= 80 and stripped[-1] not in terminal:
            stripped = stripped + "."
        result.append(stripped)
    return "\n".join(result)
def _clean_pdf_chunk_text(text: str, page_number: int | None = None) -> str:
    cleaned = _strip_pdf_page_number_lines(text, page_number)
    # Terminal punctuation injected at extraction time via font-metadata
    # heading detection in _extract_formatted_pdf. Body text never touched.
    cleaned = _sanitize_text_for_db(_reflow_pdf_text(cleaned))
    cleaned = _strip_pdf_page_number_prefix(cleaned, page_number)
    cleaned = _strip_pdf_page_number_suffix(cleaned, page_number)
    cleaned = _strip_pdf_page_number_edge_tokens(cleaned, page_number)
    cleaned = _drop_duplicate_leading_sentence(cleaned)
    return _sanitize_text_for_db(cleaned).strip()
def _build_sentence_boundaries(text: str) -> list[list[int]]:
    if not text:
        return []
    segmenter = pysbd.Segmenter(language="en", clean=False)
    boundaries: list[list[int]] = []
    line_offset = 0
    for line in text.splitlines(keepends=True):
        line_text = line.rstrip("\r\n")
        if not line_text.strip():
            line_offset += len(line)
            continue
        sent_cursor = 0
        for sent in segmenter.segment(line_text):
            if not sent or not sent.strip():
                continue
            try:
                rel_start = line_text.index(sent, sent_cursor)
            except ValueError:
                continue
            rel_end = rel_start + len(sent)
            boundaries.append([line_offset + rel_start, line_offset + rel_end])
            sent_cursor = rel_end
        line_offset += len(line)
    return boundaries
router = APIRouter()

ALLOWED_EXTENSIONS = frozenset({".pdf", ".docx", ".txt", ".md", ".html", ".epub"})
TEXT_EXTENSIONS = frozenset({".txt", ".md", ".html"})
TARGET_CHUNK_SIZE = 1500  # characters per chunk

def _extract_formatted_docx(file_bytes: bytes) -> tuple[str, list[dict]] | None:
    """Extract plain text AND formatted_content from a DOCX file.

    Returns (plain_text, formatted_content_list) or None on failure.
    """
    try:
        import docx  # python-docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # M13.4 Ship 1: reverse map for paragraph alignment so uploaded
        # DOCX files preserve the user's center/right/justify intent
        # through to the Quill editor and back to the downloaded export.
        # Without this, R9 from the diagnostic would silently strip
        # alignment at upload time.
        _ALIGN_REVERSE = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        }

        doc = docx.Document(io.BytesIO(file_bytes))
        formatted: list[dict] = []
        plain_parts: list[str] = []

        for para in doc.paragraphs:
            para_text = (para.text or "").strip()
            if not para_text:
                continue

            # Determine paragraph type
            style_name = (para.style.name or "").lower() if para.style else ""
            ptype = "paragraph"
            level = None

            if style_name.startswith("heading"):
                ptype = "heading"
                # Extract level from style name like "Heading 1", "Heading 2"
                for ch in style_name:
                    if ch.isdigit():
                        level = int(ch)
                        break
                if level is None:
                    level = 1
            elif style_name.startswith("list") or style_name.startswith("bullet"):
                ptype = "list_item"

            # Extract runs with formatting
            runs: list[dict] = []
            for run in para.runs:
                run_text = run.text or ""
                if not run_text:
                    continue
                run_data: dict = {"text": run_text}
                if run.bold:
                    run_data["bold"] = True
                if run.italic:
                    run_data["italic"] = True
                if run.underline:
                    run_data["underline"] = True
                if run.font and run.font.strike:
                    run_data["strike"] = True
                if run.font and run.font.size:
                    run_data["font_size"] = round(run.font.size.pt, 1)
                # Color: python-docx returns RGBColor; serialize to
                # lowercase 6-digit hex without `#` to match the canonical
                # storage shape from the save path normalizer.
                if (
                    run.font
                    and run.font.color
                    and run.font.color.rgb is not None
                ):
                    rgb = run.font.color.rgb
                    try:
                        run_data["color"] = str(rgb).lower()
                    except (TypeError, ValueError):
                        pass
                # Font family: run.font.name returns the ascii face if
                # set on the run. None means "inherit from style cascade"
                # — don't emit so the Word style decides at re-render.
                if run.font and run.font.name:
                    run_data["font_family"] = run.font.name
                runs.append(run_data)

            # If no runs extracted (e.g. field codes), fall back to full text
            if not runs:
                runs = [{"text": para_text}]

            entry: dict = {"type": ptype, "runs": runs}
            if level is not None:
                entry["level"] = level
            # Block-level alignment (M13.4 Ship 1). para.alignment is
            # None when the paragraph inherits from the style cascade —
            # we don't emit in that case so the Quill editor and the
            # downstream export both honor the cascade.
            align_canon = _ALIGN_REVERSE.get(para.alignment)
            if align_canon is not None:
                entry["alignment"] = align_canon
            formatted.append(entry)
            plain_parts.append(para_text)

        plain = "\n\n".join(plain_parts).strip()
        return (plain, formatted) if plain else None
    except Exception:
        logger.exception("docx.format_extract.failed")
        return None


def _extract_formatted_pdf(file_bytes: bytes) -> tuple[list[dict], list[list[dict]]] | None:
    """Extract page text AND per-page formatted_content from a PDF.

    Uses pdfplumber to access per-character font metadata so headings can be
    identified by size/weight rather than text-case heuristics.  Terminal
    punctuation is injected ONLY on heading lines that lack it, giving TTS a
    pause cue without ever modifying body text.

    Returns (pages_list, per_page_formatted) or None on failure.
    pages_list items: {"page_number": int, "text": str}
    per_page_formatted: list of formatted blocks per page.
    """
    try:
        import pdfplumber  # noqa: F811
    except ImportError:
        logger.error("pdf.format_extract.pdfplumber_not_installed")
        return None

    _TERMINAL_RE = re.compile(r"[.!?:;…]\s*$")
    _BUCKET = 2  # pts — chars within this y-range are on the same visual line

    try:
        pdf = pdfplumber.open(io.BytesIO(file_bytes))
        pages: list[dict] = []
        per_page_formatted: list[list[dict]] = []

        for page_idx, page in enumerate(pdf.pages):
            chars = page.chars
            if not chars:
                continue

            # --- group chars into visual lines by y0 bucket ---------------
            lines_by_y: dict[float, list[dict]] = {}
            for ch in chars:
                bucket = round(ch["top"] / _BUCKET) * _BUCKET
                lines_by_y.setdefault(bucket, []).append(ch)

            # sort buckets top-to-bottom, chars left-to-right within a line
            sorted_buckets = sorted(lines_by_y.keys())
            visual_lines: list[tuple[str, float, bool]] = []  # (text, mean_size, is_bold)
            all_sizes: list[float] = []

            for bucket in sorted_buckets:
                line_chars = sorted(lines_by_y[bucket], key=lambda c: c["x0"])
                text = "".join(c["text"] for c in line_chars).strip()
                if not text:
                    continue
                sizes = [c.get("size", 0) for c in line_chars if c["text"].strip()]
                mean_size = sum(sizes) / len(sizes) if sizes else 0
                is_bold = any(
                    "bold" in (c.get("fontname", "") or "").lower()
                    for c in line_chars
                    if c["text"].strip()
                )
                visual_lines.append((text, mean_size, is_bold))
                all_sizes.extend(sizes)

            if not visual_lines:
                continue

            # --- compute body font-size median ----------------------------
            all_sizes.sort()
            mid = len(all_sizes) // 2
            if len(all_sizes) % 2 == 1:
                body_median = all_sizes[mid]
            else:
                body_median = (all_sizes[mid - 1] + all_sizes[mid]) / 2 if all_sizes else 0

            # --- classify lines & build text + blocks ---------------------
            text_parts: list[str] = []
            page_blocks: list[dict] = []
            prev_was_heading = False

            for line_text, mean_size, is_bold in visual_lines:
                is_heading = False
                level = None

                if body_median > 0:
                    if mean_size >= body_median * 1.15:
                        is_heading = True
                        level = 1
                    elif is_bold and mean_size >= body_median:
                        is_heading = True
                        level = 2

                if is_heading:
                    # Inject terminal period on headings missing punctuation
                    if not _TERMINAL_RE.search(line_text):
                        line_text = line_text + "."
                    # Blank line before heading for paragraph structure
                    if text_parts and text_parts[-1] != "":
                        text_parts.append("")
                    text_parts.append(line_text)
                    text_parts.append("")  # blank line after heading
                    prev_was_heading = True

                    entry: dict = {"type": "heading", "level": level, "runs": [{"text": line_text}]}
                    page_blocks.append(entry)
                else:
                    text_parts.append(line_text)
                    prev_was_heading = False

                    entry = {"type": "paragraph", "runs": [{"text": line_text}]}
                    page_blocks.append(entry)

            raw_text = "\n".join(text_parts).strip()
            clean = _clean_pdf_chunk_text(raw_text, page_idx + 1)
            if not clean.strip():
                continue

            pages.append({"page_number": page_idx + 1, "text": clean})
            per_page_formatted.append(page_blocks)

        pdf.close()
        return (pages, per_page_formatted) if pages else None
    except Exception:
        logger.exception("pdf.format_extract.failed")
        return None


def _extract_text(file_bytes: bytes, extension: str) -> str | list[dict] | None:
    """Extract plain text from supported file types."""
    if extension in TEXT_EXTENSIONS:
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue

    if extension == ".docx":
        result = _extract_formatted_docx(file_bytes)
        if result:
            return result[0]  # plain text only for legacy path
        return None

    if extension == ".pdf":
        result = _extract_formatted_pdf(file_bytes)
        if result:
            return result[0]  # pages list only for legacy path
        return None

    if extension == ".epub":
        return _extract_epub_text(file_bytes)

    return None


def _extract_epub_text(file_bytes: bytes) -> str | None:
    """Extract readable text from an EPUB.

    An EPUB is a ZIP of XHTML chapters plus an OPF package file that defines the
    reading order (spine). We resolve the spine, strip each chapter's markup to
    text, and join the chapters in order. Pure stdlib (zipfile + html.parser) —
    no extra dependency. Returns None on anything unparseable so the caller logs
    "unsupported" and fails gracefully.
    """
    import posixpath
    import zipfile
    import xml.etree.ElementTree as ET
    from html.parser import HTMLParser

    _BLOCK_TAGS = {
        "p", "div", "br", "li", "tr", "section", "article",
        "h1", "h2", "h3", "h4", "h5", "h6", "blockquote",
    }

    class _HtmlToText(HTMLParser):
        def __init__(self) -> None:
            super().__init__(convert_charrefs=True)
            self._parts: list[str] = []
            self._skip = 0

        def handle_starttag(self, tag: str, attrs: object) -> None:
            if tag in ("script", "style"):
                self._skip += 1
            elif tag in _BLOCK_TAGS:
                self._parts.append("\n")

        def handle_endtag(self, tag: str) -> None:
            if tag in ("script", "style") and self._skip:
                self._skip -= 1
            elif tag in _BLOCK_TAGS:
                self._parts.append("\n")

        def handle_data(self, data: str) -> None:
            if self._skip == 0 and data.strip():
                self._parts.append(data)

        def text(self) -> str:
            return "".join(self._parts)

    def _html_to_text(raw: bytes) -> str:
        s = None
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                s = raw.decode(enc)
                break
            except UnicodeDecodeError:
                continue
        if s is None:
            return ""
        try:
            parser = _HtmlToText()
            parser.feed(s)
            return parser.text()
        except Exception:
            return re.sub(r"<[^>]+>", " ", s)

    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            names = set(zf.namelist())

            # Locate the OPF package file (preferred via META-INF/container.xml).
            opf_path: str | None = None
            if "META-INF/container.xml" in names:
                try:
                    root = ET.fromstring(zf.read("META-INF/container.xml"))
                    for el in root.iter():
                        if el.tag.endswith("rootfile") and el.get("full-path"):
                            opf_path = el.get("full-path")
                            break
                except ET.ParseError:
                    opf_path = None
            if not opf_path:
                opf_path = next(
                    (n for n in names if n.lower().endswith(".opf")), None
                )
            if not opf_path or opf_path not in names:
                return None

            base = posixpath.dirname(opf_path)
            opf = ET.fromstring(zf.read(opf_path))

            manifest: dict[str, str] = {}
            for el in opf.iter():
                if el.tag.endswith("item") and el.get("id") and el.get("href"):
                    manifest[el.get("id")] = el.get("href")
            spine_ids = [
                el.get("idref")
                for el in opf.iter()
                if el.tag.endswith("itemref") and el.get("idref")
            ]
            hrefs = [manifest[i] for i in spine_ids if i in manifest]
            if not hrefs:
                hrefs = list(manifest.values())

            parts: list[str] = []
            for href in hrefs:
                href = href.split("#", 1)[0]
                candidate = (
                    posixpath.normpath(posixpath.join(base, href))
                    if base
                    else href
                )
                path = candidate if candidate in names else (
                    href if href in names else None
                )
                if path is None or not path.lower().endswith(
                    (".xhtml", ".html", ".htm")
                ):
                    continue
                chapter = _html_to_text(zf.read(path)).strip()
                if chapter:
                    parts.append(chapter)

            text = "\n\n".join(parts).strip()
            return text or None
    except Exception:
        return None


def _chunk_markdown(raw_text: str) -> list[dict]:
    """Split text into chunks by headings or paragraph groups."""
    lines = raw_text.strip().splitlines()
    if not lines:
        return []

    sections: list[tuple[str, list[str]]] = []
    current_title = "Section 1"
    current_lines: list[str] = []

    for line in lines:
        heading_match = re.match(r"^(#{1,3})\s+(.+)", line)
        if heading_match:
            if current_lines:
                sections.append((current_title, current_lines))
                current_lines = []
            current_title = heading_match.group(2).strip()
        else:
            current_lines.append(line)

    if current_lines:
        sections.append((current_title, current_lines))

    if not sections:
        return []

    chunks: list[dict] = []
    seq = 0

    for title, sec_lines in sections:
        body = "\n".join(sec_lines).strip()
        if not body:
            continue

        if len(body) <= TARGET_CHUNK_SIZE * 1.5:
            chunks.append({"title": title, "text": body, "seq": seq})
            seq += 1
            continue

        paragraphs = re.split(r"\n\s*\n", body)
        buffer = ""
        part = 1
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            if buffer and len(buffer) + len(para) > TARGET_CHUNK_SIZE:
                chunk_title = title if part == 1 else f"{title} (part {part})"
                chunks.append({"title": chunk_title, "text": buffer.strip(), "seq": seq})
                seq += 1
                part += 1
                buffer = para + "\n\n"
            else:
                buffer += para + "\n\n"

        if buffer.strip():
            chunk_title = title if part == 1 else f"{title} (part {part})"
            chunks.append({"title": chunk_title, "text": buffer.strip(), "seq": seq})
            seq += 1

    if not chunks and raw_text.strip():
        chunks.append({"title": "Document", "text": raw_text.strip()[:5000], "seq": 0})

    return chunks


def _chunk_by_pages(pages: list[dict]) -> list[dict]:
    """Split PDF pages into chunks — one page = one chunk.

    If a page exceeds 5000 characters it is split into sub-chunks.
    Each chunk dict has: title, text, seq, page_number.
    """
    chunks: list[dict] = []
    seq = 0
    max_chars = 5000

    for page in pages:
        page_num = page["page_number"]
        page_text = page["text"]

        if len(page_text) <= max_chars:
            chunks.append({
                "title": f"Page {page_num}",
                "text": page_text,
                "seq": seq,
                "page_number": page_num,
            })
            seq += 1
        else:
            # Split oversized page on paragraph boundaries
            paragraphs = re.split(r"\n\s*\n", page_text)
            buffer = ""
            part = 1
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                if buffer and len(buffer) + len(para) > max_chars:
                    chunks.append({
                        "title": f"Page {page_num} (part {part})",
                        "text": buffer.strip(),
                        "seq": seq,
                        "page_number": page_num,
                    })
                    seq += 1
                    part += 1
                    buffer = para + "\n\n"
                else:
                    buffer += para + "\n\n"
            if buffer.strip():
                chunks.append({
                    "title": f"Page {page_num} (part {part})" if part > 1 else f"Page {page_num}",
                    "text": buffer.strip(),
                    "seq": seq,
                    "page_number": page_num,
                })
                seq += 1

    return chunks


async def _process_document(
    doc_id: UUID,
    file_bytes: bytes,
    extension: str,
    db: AsyncSession,
    *,
    page_texts_json: str | None = None,
) -> tuple[int, list[UUID]]:
    """Extract text, chunk, and insert into document_chunks. Returns (chunk_count, chunk_ids)."""

    # ── Extract formatted content alongside plain text ──────────────────
    chunk_ids: list[UUID] = []
    formatted_by_chunk: dict[int, list[dict]] = {}  # seq -> formatted blocks

    if extension == ".docx":
        docx_result = _extract_formatted_docx(file_bytes)
        if docx_result is None:
            logger.warning("document.process.unsupported", doc_id=str(doc_id), ext=extension)
            return 0, []
        plain_text, all_formatted = docx_result
        raw_text = _sanitize_text_for_db(plain_text)
        chunks = _chunk_markdown(raw_text)
        page_count = len(chunks)
        all_text = raw_text

        # Map formatted blocks to chunks by matching text content
        # Each chunk gets the formatted blocks whose text falls within it
        fmt_cursor = 0
        for chunk in chunks:
            chunk_blocks: list[dict] = []
            chunk_plain = chunk["text"]
            chars_remaining = len(chunk_plain)
            while fmt_cursor < len(all_formatted) and chars_remaining > 0:
                block = all_formatted[fmt_cursor]
                block_text = "".join(r["text"] for r in block["runs"])
                chars_remaining -= len(block_text) + 2  # +2 for \n\n separator
                chunk_blocks.append(block)
                fmt_cursor += 1
            formatted_by_chunk[chunk["seq"]] = chunk_blocks

    elif extension == ".pdf":
        pages: list[dict] | None = None
        per_page_formatted: list[list[dict]] | None = None

        # --- Try client-supplied page texts (pdfrx / "pdium" source) ------
        if page_texts_json:
            try:
                raw_pages = json.loads(page_texts_json)
                cleaned_pages: list[dict] = []
                cleaned_fmt: list[list[dict]] = []
                for entry in raw_pages:
                    raw = entry.get("text", "")
                    if not raw or not raw.strip():
                        continue
                    clean = _clean_pdf_chunk_text(raw, entry.get("page_number"))
                    if not clean.strip():
                        continue
                    cleaned_pages.append({
                        "page_number": entry["page_number"],
                        "text": clean,
                    })
                    # Build formatted blocks: split on blank lines → paragraphs
                    paragraphs = re.split(r"\n\s*\n", clean)
                    blocks: list[dict] = []
                    for para in paragraphs:
                        para = para.strip()
                        if para:
                            blocks.append({"type": "paragraph", "runs": [{"text": para}]})
                    cleaned_fmt.append(blocks)
                if cleaned_pages:
                    pages = cleaned_pages
                    per_page_formatted = cleaned_fmt
                    logger.info(
                        "pdf.extract.pdium_source",
                        doc_id=str(doc_id),
                        page_count=len(pages),
                    )
            except (json.JSONDecodeError, Exception) as exc:
                logger.warning(
                    "pdf.extract.pdium_parse_failed",
                    doc_id=str(doc_id),
                    error=str(exc),
                )
                pages = None
                per_page_formatted = None

        # --- Fallback to server-side pdfplumber extraction ----------------
        if pages is None:
            pdf_result = _extract_formatted_pdf(file_bytes)
            if pdf_result is None:
                logger.warning("document.process.unsupported", doc_id=str(doc_id), ext=extension)
                return 0, []
            pages, per_page_formatted = pdf_result

        chunks = _chunk_by_pages(pages)
        page_count = len(pages)
        all_text = " ".join(p["text"] for p in pages)

        # Map formatted blocks by page number -> chunk seq
        page_fmt: dict[int, list[dict]] = {}
        for pg, fmt_blocks in zip(pages, per_page_formatted):
            page_fmt[pg["page_number"]] = fmt_blocks
        for chunk in chunks:
            pn = chunk.get("page_number", 1)
            if pn in page_fmt:
                formatted_by_chunk[chunk["seq"]] = page_fmt[pn]

    else:
        extracted = _extract_text(file_bytes, extension)
        if extracted is None:
            logger.warning("document.process.unsupported", doc_id=str(doc_id), ext=extension)
            return 0, []
        raw_text = _sanitize_text_for_db(extracted)
        chunks = _chunk_markdown(raw_text)
        page_count = len(chunks)
        all_text = raw_text

    if not chunks:
        return 0, []

    word_count = len([w for w in re.split(r"\s+", all_text.strip()) if w])

    insert_chunk_sql = text(
        "INSERT INTO document_chunks "
        "(id, document_id, sequence_index, chunk_type, text_content, tone, page_number, character_count, metadata_json, formatted_content) "
        "VALUES (:id, :doc_id, :seq, :ctype, :txt, :tone, :page, :chars, :meta, :fmt)"
    )

    for chunk in chunks:
        chunk_id = uuid4()
        chunk_ids.append(chunk_id)
        chunk_text = chunk["text"]
        boundaries = _build_sentence_boundaries(chunk_text)
        existing_meta = chunk.get("metadata_json") or {}
        existing_meta["sentence_boundaries"] = boundaries
        chunk["metadata_json"] = existing_meta
        meta_json = json.dumps({"title": chunk["title"], **chunk["metadata_json"]})

        fmt_content = formatted_by_chunk.get(chunk["seq"])
        fmt_json = json.dumps(fmt_content, ensure_ascii=False) if fmt_content else None

        await db.execute(
            insert_chunk_sql,
            {
                "id": chunk_id,
                "doc_id": doc_id,
                "seq": chunk["seq"],
                "ctype": "text",
                "txt": chunk["text"],
                "tone": "neutral",
                "page": chunk.get("page_number", 1),
                "chars": len(chunk["text"]),
                "meta": meta_json,
                "fmt": fmt_json,
            },
        )

    await db.execute(
        text(
            "UPDATE documents "
            "SET status = :st, page_count = :pc, word_count = :wc, updated_at = NOW() "
            "WHERE id = :did"
        ),
        {"st": "ready", "pc": page_count, "wc": word_count, "did": doc_id},
    )

    logger.info("document.processed", doc_id=str(doc_id), chunks=len(chunk_ids))
    return len(chunk_ids), chunk_ids


def _extract_epub_cover_bytes(file_bytes: bytes) -> tuple[bytes, str] | None:
    """Find and return the EPUB's embedded cover image as (bytes, extension).

    Resolution order: EPUB3 manifest item with properties="cover-image";
    EPUB2 <meta name="cover" content="id">; then any image item whose id/href
    contains "cover". Returns None if no cover image is found.
    """
    import posixpath
    import zipfile
    import xml.etree.ElementTree as ET

    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            names = set(zf.namelist())
            opf_path: str | None = None
            if "META-INF/container.xml" in names:
                try:
                    root = ET.fromstring(zf.read("META-INF/container.xml"))
                    for el in root.iter():
                        if el.tag.endswith("rootfile") and el.get("full-path"):
                            opf_path = el.get("full-path")
                            break
                except ET.ParseError:
                    opf_path = None
            if not opf_path:
                opf_path = next(
                    (n for n in names if n.lower().endswith(".opf")), None
                )
            if not opf_path or opf_path not in names:
                return None

            base = posixpath.dirname(opf_path)
            opf = ET.fromstring(zf.read(opf_path))

            # manifest id -> (href, media_type, properties)
            manifest: dict[str, tuple[str, str, str]] = {}
            for el in opf.iter():
                if el.tag.endswith("item") and el.get("id") and el.get("href"):
                    manifest[el.get("id")] = (
                        el.get("href"),
                        (el.get("media-type") or "").lower(),
                        (el.get("properties") or "").lower(),
                    )

            cover_href: str | None = None
            for href, _mt, props in manifest.values():
                if "cover-image" in props:
                    cover_href = href
                    break
            if not cover_href:
                cover_id = None
                for el in opf.iter():
                    if (
                        el.tag.endswith("meta")
                        and (el.get("name") or "").lower() == "cover"
                        and el.get("content")
                    ):
                        cover_id = el.get("content")
                        break
                if cover_id and cover_id in manifest:
                    cover_href = manifest[cover_id][0]
            if not cover_href:
                for cid, (href, mt, _props) in manifest.items():
                    if mt.startswith("image/") and (
                        "cover" in cid.lower() or "cover" in href.lower()
                    ):
                        cover_href = href
                        break
            if not cover_href:
                return None

            cover_href = cover_href.split("#", 1)[0]
            candidate = (
                posixpath.normpath(posixpath.join(base, cover_href))
                if base
                else cover_href
            )
            path = candidate if candidate in names else (
                cover_href if cover_href in names else None
            )
            if not path:
                return None
            ext = path.rsplit(".", 1)[-1].lower() if "." in path else "jpg"
            return zf.read(path), ext
    except Exception:
        return None


async def _autostore_epub_cover(
    doc_id: UUID, file_bytes: bytes, db: AsyncSession
) -> tuple[str, str] | None:
    """Extract the EPUB's cover image, resize to a 1600px JPEG, store it under
    the document's cover key, and set cover_type/cover_value. Fully defensive —
    any failure leaves the document without a cover (it falls back to the type
    banner). Returns (cover_type, cover_value) or None.
    """
    try:
        found = _extract_epub_cover_bytes(file_bytes)
        if found is None:
            return None
        img_bytes, _src_ext = found

        from PIL import Image

        img = Image.open(io.BytesIO(img_bytes))
        if img.mode in ("RGBA", "P", "LA"):
            img = img.convert("RGB")
        max_dim = 1600
        if img.width > max_dim or img.height > max_dim:
            img.thumbnail((max_dim, max_dim), Image.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=88)
        out_bytes = buf.getvalue()

        settings = get_settings()
        s3 = S3StorageProvider(settings)
        bucket = settings.S3_BUCKET_NAME
        key = f"covers/{doc_id}.jpg"
        try:
            await s3.delete_by_prefix(bucket, f"covers/{doc_id}.")
        except Exception:
            pass
        await s3.put_object(bucket, key, out_bytes, content_type="image/jpeg")

        await db.execute(
            text(
                "UPDATE documents SET cover_type = 'uploaded', cover_value = :cv "
                "WHERE id = :id"
            ),
            {"cv": key, "id": str(doc_id)},
        )
        logger.info("document.epub.cover_autostored", doc_id=str(doc_id))
        return "uploaded", key
    except Exception:
        logger.warning("document.epub.cover_autostore_failed", doc_id=str(doc_id))
        return None


@router.post("/blank/", status_code=status.HTTP_201_CREATED)
async def create_blank_document(
    request: Request,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> dict:
    """Create a new blank DOCX document.

    source_type='docx' enables the full DOCX editor pipeline
    (DocxDocumentViewport + DocxDocumentEditor) on the client, giving users
    a Word-like page-styled editor with toolbar out of the box. Renamed
    from the legacy 'blank' source type which used the inline plain-text
    chunk editor.

    The chunk is seeded with an empty paragraph block matching the
    formatted_content schema produced by _extract_formatted_docx (JSONB
    list of {"type": "paragraph"|"heading"|"list_item", "runs": [...]}).
    """
    from psitta.services.subscription_service import check_and_increment_doc_quota

    await check_and_increment_doc_quota(db, user_id)

    doc_id = uuid4()
    chunk_id = uuid4()
    now = datetime.now(timezone.utc)

    # One empty paragraph block — minimum valid DOCX block structure.
    empty_docx_blocks = [{"type": "paragraph", "runs": [{"text": ""}]}]

    await db.execute(
        text(
            "INSERT INTO documents "
            "(id, user_id, title, source_type, status, file_size_bytes, storage_key, page_count, word_count, created_at, updated_at) "
            "VALUES "
            "(:id, :user_id, :title, :source_type, :status, 0, '', 1, 0, :now, :now)"
        ),
        {
            "id": doc_id,
            "user_id": str(user_id),
            "title": "Untitled Sheet",
            "source_type": "docx",
            "status": "ready",
            "now": now,
        },
    )

    await db.execute(
        text(
            "INSERT INTO document_chunks "
            "(id, document_id, sequence_index, text_content, character_count, created_at, formatted_content) "
            "VALUES "
            "(:id, :doc_id, 0, '', 0, :now, CAST(:fmt AS jsonb))"
        ),
        {
            "id": chunk_id,
            "doc_id": doc_id,
            "now": now,
            "fmt": json.dumps(empty_docx_blocks),
        },
    )

    logger.info("document.blank.created", doc_id=str(doc_id), chunk_id=str(chunk_id))

    await audit_service.log_event(
        db,
        action="document.create_blank",
        resource_type="document",
        user_id=str(user_id),
        resource_id=str(doc_id),
        details={"title": "Untitled Sheet"},
        ip_address=request.client.host if request.client else None,
    )
    return {
        "id": str(doc_id),
        "chunk_id": str(chunk_id),
        "title": "Untitled Sheet",
        "status": "ready",
        "source_type": "docx",
        "created_at": now.isoformat(),
    }


@router.post("/", status_code=status.HTTP_202_ACCEPTED)
async def upload_document(
    file: UploadFile,
    background_tasks: BackgroundTasks,
    request: Request,
    page_texts: str | None = Form(None),
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> dict:
    filename = file.filename or "unknown"
    extension = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    logger.info(
        "upload.trace.documents.upload.begin",
        filename=filename,
        extension=extension,
        user_id=str(user_id),
    )

    if extension not in ALLOWED_EXTENSIONS:
        logger.warning(
            "upload.trace.documents.upload.unsupported_extension",
            filename=filename,
            extension=extension,
            allowed_extensions=sorted(ALLOWED_EXTENSIONS),
        )
        raise HTTPException(status_code=415, detail="Unsupported file type")

    file_bytes = await file.read()
    file_size = len(file_bytes)

    doc_id = uuid4()
    logger.info(
        "upload.trace.documents.upload.file_read",
        filename=filename,
        extension=extension,
        file_size_bytes=file_size,
        doc_id=str(doc_id),
    )

    from psitta.services.audio_cache import put_raw_file
    try:
        logger.info(
            "upload.trace.documents.raw_persist.begin",
            doc_id=str(doc_id),
            filename=filename,
            extension=extension,
            file_size_bytes=file_size,
        )
        await put_raw_file(str(doc_id), extension, file_bytes)
        logger.info(
            "upload.trace.documents.raw_persist.end",
            doc_id=str(doc_id),
            filename=filename,
            extension=extension,
            file_size_bytes=file_size,
        )
    except Exception as e:
        logger.error(
            "upload.trace.documents.upload.abort",
            doc_id=str(doc_id),
            filename=filename,
            extension=extension,
            exception_type=type(e).__name__,
            exception_message=str(e),
        )
        logger.exception(
            "document.upload.raw_persist_failed",
            doc_id=str(doc_id),
            filename=filename,
            error=str(e),
        )
        raise HTTPException(
            status_code=503,
            detail="Original file storage unavailable. Upload aborted.",
        ) from e

    from psitta.services.subscription_service import check_and_increment_doc_quota
    await check_and_increment_doc_quota(db, user_id)

    title = filename.rsplit(".", 1)[0] if "." in filename else filename
    source_type = extension.lstrip(".")
    storage_key = f"uploads/{doc_id}{extension}"
    now = datetime.now(timezone.utc)

    await db.execute(
        text(
            "INSERT INTO documents "
            "(id, user_id, title, source_type, status, file_size_bytes, storage_key, page_count, word_count, created_at, updated_at) "
            "VALUES "
            "(:id, :user_id, :title, :source_type, :status, :file_size_bytes, :storage_key, :page_count, :word_count, :created_at, :updated_at)"
        ),
        {
            "id": doc_id,
            "user_id": str(user_id),
            "title": title,
            "source_type": source_type,
            "status": "uploaded",
            "file_size_bytes": file_size,
            "storage_key": storage_key,
            "page_count": 0,
            "word_count": 0,
            "created_at": now,
            "updated_at": now,
        },
    )

    chunk_count, chunk_ids = await _process_document(doc_id, file_bytes, extension, db, page_texts_json=page_texts)
    doc_status = "ready" if chunk_count > 0 else "uploaded"
    if chunk_count > 0:
        background_tasks.add_task(_eager_synthesize_chunks, doc_id, chunk_ids, user_id)
        logger.info("tts.eager_synthesis.queued", doc_id=str(doc_id), chunks=chunk_count)

    # EPUBs ship with their own cover art — adopt it automatically so the
    # uploaded book looks polished (writer can still change it in the Library).
    auto_cover_type: str | None = None
    auto_cover_value: str | None = None
    if extension == ".epub":
        auto = await _autostore_epub_cover(doc_id, file_bytes, db)
        if auto is not None:
            auto_cover_type, auto_cover_value = auto

    logger.info("document.upload.accepted", doc_id=str(doc_id), title=title, chunks=chunk_count)
    logger.info(
        "upload.trace.documents.upload.complete",
        doc_id=str(doc_id),
        title=title,
        source_type=source_type,
        status=doc_status,
        chunk_count=chunk_count,
    )

    await audit_service.log_event(
        db,
        action="document.upload",
        resource_type="document",
        user_id=str(user_id),
        resource_id=str(doc_id),
        details={
            "filename": filename,
            "source_type": source_type,
            "file_size_bytes": file_size,
            "chunk_count": chunk_count,
        },
        ip_address=request.client.host if request.client else None,
    )
    return {
        "id": str(doc_id),
        "title": title,
        "status": doc_status,
        "source_type": source_type,
        "page_count": chunk_count,
        "created_at": now.isoformat(),
        "cover_type": auto_cover_type,
        "cover_value": auto_cover_value,
    }


async def _maybe_seed_welcome_kit(
    db: AsyncSession,
    user_id: UUID,
    background_tasks: BackgroundTasks,
) -> None:
    """Claim and schedule one-time welcome-kit seeding for Writing Nook writers.

    Cheap fast-path: a single boolean SELECT for already-seeded users (the
    common case). Only unseeded users pay for plan resolution. The claim
    UPDATE is atomic (``WHERE welcome_seeded = false RETURNING id``) so
    concurrent Library loads seed at most once.

    Fully fail-open AND transaction-safe: all DB work runs inside a
    SAVEPOINT (``begin_nested``). If anything fails — including the
    ``welcome_seeded`` column not existing yet because the migration hasn't
    been applied — the savepoint rolls back and the outer request
    transaction stays usable, so the document list query that follows still
    succeeds. A bug here must never 500 the Library.
    """
    schedule = False
    try:
        async with db.begin_nested():
            already = (
                await db.execute(
                    text("SELECT welcome_seeded FROM users WHERE id = :uid"),
                    {"uid": str(user_id)},
                )
            ).scalar()
            if already is False:  # not yet seeded (True/None ⇒ skip)
                from psitta.services.subscription_service import (
                    get_effective_plan,
                )

                plan = await get_effective_plan(db, user_id)
                if plan.plan_id == "writing_nook_pro":
                    claimed = (
                        await db.execute(
                            text(
                                "UPDATE users SET welcome_seeded = true "
                                "WHERE id = :uid AND welcome_seeded = false "
                                "RETURNING id"
                            ),
                            {"uid": str(user_id)},
                        )
                    ).first()
                    schedule = claimed is not None
    except Exception:
        logger.warning(
            "welcome_seed.trigger_failed", user_id=str(user_id), exc_info=True
        )
        return

    if schedule:
        from psitta.services.seed_service import run_welcome_seed_background

        background_tasks.add_task(run_welcome_seed_background, user_id)
        logger.info("welcome_seed.scheduled", user_id=str(user_id))


@router.get("/")
async def list_documents(
    background_tasks: BackgroundTasks,
    page: Annotated[int, Query(ge=1)] = 1,
    size: Annotated[int, Query(ge=1, le=100)] = 20,
    show_archived: bool = False,
    trashed: bool = False,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> dict:
    offset = (page - 1) * size

    _cols = (
        "id, title, status, source_type, page_count, word_count, created_at, "
        "updated_at, project_id, cover_type, cover_value"
    )
    if trashed:
        # Trash view: only soft-deleted docs. No seeding here.
        where = "user_id = :uid AND status = 'deleted'"
        params: dict = {"uid": str(user_id)}
    else:
        # ── Welcome-kit seeding (Writing Nook, first login) ──────────────
        # On a writer's first Library load, seed the 6 starter documents once.
        # Fail-open: any error here must never break the document list.
        await _maybe_seed_welcome_kit(db, user_id, background_tasks)
        where = (
            "user_id = :uid AND status != 'deleted' "
            "AND (:show_archived OR status != 'archived')"
        )
        params = {"uid": str(user_id), "show_archived": show_archived}

    count_result = await db.execute(
        text(f"SELECT COUNT(*) FROM documents WHERE {where}"),  # noqa: S608 — where is a fixed literal
        params,
    )
    total = count_result.scalar() or 0

    rows = await db.execute(
        text(
            f"SELECT {_cols} FROM documents WHERE {where} "  # noqa: S608 — fixed literals, params bound
            "ORDER BY created_at DESC LIMIT :lim OFFSET :off"
        ),
        {**params, "lim": size, "off": offset},
    )

    items = [
        {
            "id": str(r.id),
            "title": r.title,
            "status": r.status,
            "source_type": r.source_type,
            "page_count": r.page_count,
            "word_count": getattr(r, "word_count", 0),
            "created_at": r.created_at.isoformat() if r.created_at else None,
            "updated_at": r.updated_at.isoformat()
            if getattr(r, "updated_at", None)
            else None,
            "project_id": str(r.project_id) if r.project_id else None,
            "cover_type": r.cover_type,
            "cover_value": r.cover_value,
        }
        for r in rows
    ]

    return {"items": items, "page": page, "size": size, "total": total}


# NOTE: this static path MUST stay above @router.get("/{document_id}") or the
# UUID route would capture "storage" and 422.
@router.get("/storage")
async def storage_usage(
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> dict:
    """Total storage used by the user's non-deleted documents."""
    row = (
        await db.execute(
            text(
                "SELECT COALESCE(SUM(file_size_bytes), 0) AS used, "
                "COUNT(*) AS cnt FROM documents "
                "WHERE user_id = :uid AND status != 'deleted'"
            ),
            {"uid": str(user_id)},
        )
    ).first()
    return {
        "used_bytes": int(row.used or 0),
        "doc_count": int(row.cnt or 0),
    }


@router.get("/{document_id}")
async def get_document(
    document_id: UUID,
    request: Request,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> dict:
    result = await db.execute(
        text(
            "SELECT id, title, status, source_type, page_count, word_count, file_size_bytes, created_at, cover_type, cover_value "
            "FROM documents WHERE id = :did AND user_id = :uid AND status != 'deleted'"
        ),
        {"did": document_id, "uid": str(user_id)},
    )
    row = result.first()
    outcome = "found" if row else "not_found_or_unauthorized"
    await audit_service.log_event(
        db,
        action="document.fetched",
        resource_type="document",
        user_id=str(user_id),
        resource_id=str(document_id),
        details={"outcome": outcome},
        ip_address=request.client.host if request.client else None,
    )
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")

    return {
        "id": str(row.id),
        "title": row.title,
        "status": row.status,
        "source_type": row.source_type,
        "page_count": row.page_count,
        "word_count": getattr(row, "word_count", 0),
        "file_size_bytes": row.file_size_bytes,
        "created_at": row.created_at.isoformat() if row.created_at else None,
        "cover_type": row.cover_type,
        "cover_value": row.cover_value,
    }


@router.get("/{document_id}/chunks")
async def get_document_chunks(
    document_id: UUID,
    request: Request,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> dict:
    """Return all chunks for a document, ordered by sequence."""
    doc_result = await db.execute(
        text(
            "SELECT id, status, source_type, chunk_positions FROM documents "
            "WHERE id = :did AND user_id = :uid AND status != 'deleted'"
        ),
        {"did": document_id, "uid": str(user_id)},
    )
    doc = doc_result.first()
    outcome = "found" if doc else "not_found_or_unauthorized"
    await audit_service.log_event(
        db,
        action="document.chunks_fetched",
        resource_type="document",
        user_id=str(user_id),
        resource_id=str(document_id),
        details={"outcome": outcome},
        ip_address=request.client.host if request.client else None,
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    rows = await db.execute(
        text(
            "SELECT id, sequence_index, chunk_type, text_content, tone, page_number, character_count, metadata_json, "
            "is_edited, edited_at, original_text, formatted_content "
            "FROM document_chunks WHERE document_id = :did "
            "ORDER BY sequence_index"
        ),
        {"did": document_id},
    )

    chunks = []
    is_pdf_document = (getattr(doc, "source_type", "") or "").lower() == "pdf"
    for r in rows:
        meta = r.metadata_json if isinstance(r.metadata_json, dict) else {}
        text_content = r.text_content
        sentence_boundaries = meta.get("sentence_boundaries")
        character_count = r.character_count
        if is_pdf_document:
            text_content = _clean_pdf_chunk_text(r.text_content or "", r.page_number)
            sentence_boundaries = _build_sentence_boundaries(text_content)
            character_count = len(text_content)
        chunks.append(
            {
                "id": str(r.id),
                "sequence_index": r.sequence_index,
                "chunk_type": r.chunk_type,
                "title": meta.get("title", f"Section {r.sequence_index + 1}"),
                "text_content": text_content,
                "tone": r.tone,
                "page_number": r.page_number,
                "character_count": character_count,
                "is_edited": getattr(r, "is_edited", False),
                "edited_at": r.edited_at.isoformat() if getattr(r, "edited_at", None) else None,
                "original_text": getattr(r, "original_text", None),
                "sentence_boundaries": sentence_boundaries,
                "formatted_content": r.formatted_content if hasattr(r, "formatted_content") else None,
            }
        )

    return {
        "document_id": str(document_id),
        "status": doc.status,
        "total_chunks": len(chunks),
        "chunks": chunks,
        # M13.1b: null for pre-M13.1b documents; client falls back to
        # computing from chunkMap in that case (lazy migration).
        "chunk_positions": doc.chunk_positions if hasattr(doc, "chunk_positions") else None,
    }


@router.get("/{document_id}/chunks/{chunk_id}/audio")
async def get_chunk_audio(
    document_id: UUID,
    chunk_id: UUID,
    request: Request,
    voice_id: str = "21m00Tcm4TlvDq8ikWAM",
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> None:
    """Stream audio for a specific chunk. Auto-synthesizes on cache miss."""
    from fastapi.responses import FileResponse

    from psitta.services.audio_cache import get_mp3, put_alignment, put_mp3, s3_key_mp3

    # Get chunk text — scoped to the authenticated user's documents.
    chunk_result = await db.execute(
        text(
            "SELECT c.text_content, c.page_number, d.source_type "
            "FROM document_chunks c "
            "JOIN documents d ON d.id = c.document_id "
            "WHERE c.id = :cid AND c.document_id = :did "
            "AND d.user_id = :uid AND d.status != 'deleted'"
        ),
        {"cid": chunk_id, "did": document_id, "uid": str(user_id)},
    )
    chunk_row = chunk_result.first()
    outcome = "found" if (chunk_row and chunk_row.text_content) else "not_found_or_unauthorized"
    await audit_service.log_event(
        db,
        action="chunk.audio_fetched",
        resource_type="document_chunk",
        user_id=str(user_id),
        resource_id=str(chunk_id),
        details={"document_id": str(document_id), "voice_id": voice_id, "outcome": outcome},
        ip_address=request.client.host if request.client else None,
    )
    if not chunk_row or not chunk_row.text_content:
        raise HTTPException(status_code=404, detail="Chunk not found")

    source_type = (getattr(chunk_row, "source_type", "") or "").lower()
    chunk_text = chunk_row.text_content
    if source_type == "pdf":
        chunk_text = _clean_pdf_chunk_text(chunk_text, getattr(chunk_row, "page_number", None))

    # Cache invalidation on edit is already handled by
    # _invalidate_chunk_audio_cache, so the previous
    # `chunk_text == chunk_row.text_content` predicate was paranoid AND
    # broken: for PDFs it always failed (chunk_text is post-cleaning,
    # text_content is raw), causing every play of every PDF chunk to
    # re-synthesize and re-bill EL. Bug-EL-1 fix.
    cached = await get_mp3(str(chunk_id), voice_id)
    if cached is not None:
        logger.info(
            "audio.cache_hit",
            chunk_id=str(chunk_id),
            voice_id=voice_id,
            source_type=source_type,
            user_id=str(user_id),
        )
        return FileResponse(cached, media_type="audio/mpeg", filename=f"{chunk_id}.mp3")

    logger.info("audio.cache_miss", chunk_id=str(chunk_id), voice_id=voice_id)

    # Synthesize via the alignment-producing quota-aware path. Using
    # the alignment variant on the audio cache-miss path means a single
    # EL bill produces BOTH the mp3 and the word-level timestamps the
    # subsequent /alignment fetch needs — the alternative
    # (synthesize_with_quota here, then waiting for /alignment to call
    # EL again) double-bills the user. Bug-EL-1 (May 3 2026): a
    # 186-char chunk was charged 372 (186 audio + 186 alignment) on
    # first play.
    from psitta.providers.tts_router import TTSRouter
    tts = TTSRouter()
    try:
        audio_bytes, alignment, provider = (
            await tts.synthesize_with_alignment_and_quota(
                chunk_text, voice_id, user_id=user_id, db=db,
            )
        )
    except Exception as e:
        logger.error("audio.synthesize_failed", error=str(e), voice_id=voice_id)
        raise HTTPException(status_code=502, detail=f"TTS synthesis failed: {e}")

    # Save mp3 to local + S3
    local_path = await put_mp3(str(chunk_id), voice_id, audio_bytes)

    # Persist alignment sidecar so a subsequent /alignment fetch hits
    # the cache instead of triggering a second EL synthesis. Only
    # ElevenLabs and Edge produce alignment with valid timestamps for
    # the audio they returned (see contract on
    # synthesize_with_alignment); discard for any other provider so a
    # future Azure path can't ship mismatched timestamps against the
    # cached audio.
    if alignment is not None and provider in {"elevenlabs", "edge"}:
        await put_alignment(
            str(chunk_id),
            voice_id,
            {
                "document_id": str(document_id),
                "chunk_id": str(chunk_id),
                "voice_id": voice_id,
                "provider": provider,
                "alignment": alignment,
            },
        )
    storage_key = s3_key_mp3(str(chunk_id), voice_id)

    # Insert cache record (delete stale row first to avoid unique constraint)
    await db.execute(
        text("DELETE FROM audio_segments WHERE chunk_id = :cid AND voice_id = :vid"),
        {"cid": chunk_id, "vid": voice_id},
    )
    from uuid import uuid4 as _uuid4
    await db.execute(
        text(
            "INSERT INTO audio_segments (id, document_id, chunk_id, voice_id, speed, storage_key, duration_ms, file_size_bytes, format) "
            "VALUES (:id, :doc_id, :chunk_id, :voice_id, :speed, :key, :dur, :size, :fmt)"
        ),
        {
            "id": _uuid4(),
            "doc_id": document_id,
            "chunk_id": chunk_id,
            "voice_id": voice_id,
            "speed": 1.0,
            "key": storage_key,
            "dur": int(len(audio_bytes) / 24),
            "size": len(audio_bytes),
            "fmt": "mp3",
        },
    )
    await db.commit()
    logger.info("audio.synthesized_on_demand", chunk_id=str(chunk_id), voice_id=voice_id, size=len(audio_bytes))

    return FileResponse(local_path, media_type="audio/mpeg", filename=f"{chunk_id}.mp3")


@router.get("/{document_id}/chunks/{chunk_id}/audio/stream")
async def get_chunk_audio_stream(
    document_id: UUID,
    chunk_id: UUID,
    request: Request,
    voice_id: str = "21m00Tcm4TlvDq8ikWAM",
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
):
    """Streaming audio for the Writing Nook: pipes TTS audio to the client as
    the model generates it so playback can start within ~1s, then writes the
    cache on completion so the next play is an instant, zero-credit hit. A
    cache hit serves the file directly. The batch /audio endpoint used by the
    Reading Nook is unchanged.
    """
    from fastapi.responses import FileResponse, StreamingResponse

    from psitta.services.audio_cache import (
        get_mp3,
        put_alignment,
        put_mp3,
        s3_key_mp3,
    )
    from psitta.services.subscription_service import check_el_quota

    # Validate + fetch chunk text, scoped to the authenticated user.
    chunk_result = await db.execute(
        text(
            "SELECT c.text_content, c.page_number, d.source_type "
            "FROM document_chunks c "
            "JOIN documents d ON d.id = c.document_id "
            "WHERE c.id = :cid AND c.document_id = :did "
            "AND d.user_id = :uid AND d.status != 'deleted'"
        ),
        {"cid": chunk_id, "did": document_id, "uid": str(user_id)},
    )
    chunk_row = chunk_result.first()
    if not chunk_row or not chunk_row.text_content:
        raise HTTPException(status_code=404, detail="Chunk not found")

    source_type = (getattr(chunk_row, "source_type", "") or "").lower()
    chunk_text = chunk_row.text_content
    if source_type == "pdf":
        chunk_text = _clean_pdf_chunk_text(
            chunk_text, getattr(chunk_row, "page_number", None)
        )

    # Cache hit → serve the file directly (instant, zero credit).
    cached = await get_mp3(str(chunk_id), voice_id)
    if cached is not None:
        logger.info(
            "audio.stream.cache_hit", chunk_id=str(chunk_id), voice_id=voice_id
        )
        return FileResponse(
            cached, media_type="audio/mpeg", filename=f"{chunk_id}.mp3"
        )

    # Decide ElevenLabs eligibility up-front while the request session is alive.
    used, limit, period_start = await check_el_quota(db, user_id)
    allow_el = not (limit > 0 and used >= limit)

    await audit_service.log_event(
        db,
        action="chunk.audio_streamed",
        resource_type="document_chunk",
        user_id=str(user_id),
        resource_id=str(chunk_id),
        details={
            "document_id": str(document_id),
            "voice_id": voice_id,
            "allow_el": allow_el,
        },
        ip_address=request.client.host if request.client else None,
    )

    from psitta.providers.tts_router import TTSRouter

    tts = TTSRouter()

    async def _audio_stream():
        from uuid import uuid4 as _uuid4

        buf = bytearray()
        alignment: dict | None = None
        provider: str | None = None
        try:
            async for ev in tts.stream_with_alignment(
                chunk_text, voice_id, allow_elevenlabs=allow_el
            ):
                if ev.get("type") == "audio":
                    data = ev["data"]
                    buf.extend(data)
                    yield data
                elif ev.get("type") == "alignment":
                    alignment = ev.get("data")
                    provider = ev.get("provider")
        except Exception as e:  # noqa: BLE001 — never crash mid-stream
            logger.error(
                "audio.stream.synth_failed",
                chunk_id=str(chunk_id),
                error=str(e),
            )
            return

        if not buf:
            logger.warning("audio.stream.empty", chunk_id=str(chunk_id))
            return

        # Cache write-through + quota increment on a FRESH session — the
        # request-scoped session has been torn down by the time the generator
        # finishes streaming.
        try:
            audio_bytes = bytes(buf)
            await put_mp3(str(chunk_id), voice_id, audio_bytes)
            if alignment is not None and provider in {"elevenlabs", "edge"}:
                await put_alignment(
                    str(chunk_id),
                    voice_id,
                    {
                        "document_id": str(document_id),
                        "chunk_id": str(chunk_id),
                        "voice_id": voice_id,
                        "provider": provider,
                        "alignment": alignment,
                    },
                )
            storage_key = s3_key_mp3(str(chunk_id), voice_id)
            from psitta.db.session import async_session_factory

            async with async_session_factory() as wdb:
                await wdb.execute(
                    text(
                        "DELETE FROM audio_segments "
                        "WHERE chunk_id = :cid AND voice_id = :vid"
                    ),
                    {"cid": chunk_id, "vid": voice_id},
                )
                await wdb.execute(
                    text(
                        "INSERT INTO audio_segments "
                        "(id, document_id, chunk_id, voice_id, speed, storage_key, "
                        "duration_ms, file_size_bytes, format) "
                        "VALUES (:id, :doc_id, :chunk_id, :voice_id, :speed, :key, "
                        ":dur, :size, :fmt)"
                    ),
                    {
                        "id": _uuid4(),
                        "doc_id": document_id,
                        "chunk_id": chunk_id,
                        "voice_id": voice_id,
                        "speed": 1.0,
                        "key": storage_key,
                        "dur": int(len(audio_bytes) / 24),
                        "size": len(audio_bytes),
                        "fmt": "mp3",
                    },
                )
                if provider == "elevenlabs" and limit > 0:
                    from psitta.services.subscription_service import (
                        increment_el_chars,
                    )

                    await increment_el_chars(
                        wdb, user_id, period_start, len(chunk_text)
                    )
                await wdb.commit()
            logger.info(
                "audio.stream.cached",
                chunk_id=str(chunk_id),
                provider=provider,
                size=len(audio_bytes),
            )
        except Exception as e:  # noqa: BLE001 — cache write is best-effort
            logger.warning(
                "audio.stream.cache_write_failed",
                chunk_id=str(chunk_id),
                error=str(e),
            )

    return StreamingResponse(_audio_stream(), media_type="audio/mpeg")


@router.get("/{document_id}/chunks/{chunk_id}/alignment")
async def get_chunk_alignment(
    document_id: UUID,
    chunk_id: UUID,
    request: Request,
    voice_id: str = "21m00Tcm4TlvDq8ikWAM",
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> dict:
    """Return alignment (timing) metadata for a chunk+voice.

    This does NOT change the /audio endpoint content-type.
    Alignment is stored as a sidecar JSON file next to the cached mp3.

    Path:
      S3: audio/{chunk_id}_{voice_id}.alignment.json
      Local cache: /tmp/psitta_audio/{chunk_id}_{voice_id}.alignment.json (ephemeral)
    """
    import json

    from psitta.services.audio_cache import get_alignment, put_alignment, put_mp3, s3_key_mp3

    # Non-ElevenLabs voices cannot produce alignment data.
    # Skip synthesis entirely and cache a null-alignment payload.
    if voice_id.startswith("en-") and voice_id.endswith("Neural"):
        payload = {
            "document_id": str(document_id),
            "chunk_id": str(chunk_id),
            "voice_id": voice_id,
            "provider": "azure",
            "alignment": None,
        }
        await put_alignment(str(chunk_id), voice_id, payload)
        return payload

    # Load chunk text — scoped to the authenticated user's documents.
    chunk_result = await db.execute(
        text(
            "SELECT c.text_content, c.page_number, d.source_type "
            "FROM document_chunks c "
            "JOIN documents d ON d.id = c.document_id "
            "WHERE c.id = :cid AND c.document_id = :did "
            "AND d.user_id = :uid AND d.status != 'deleted'"
        ),
        {"cid": chunk_id, "did": document_id, "uid": str(user_id)},
    )
    chunk_row = chunk_result.first()
    outcome = "found" if (chunk_row and chunk_row.text_content) else "not_found_or_unauthorized"
    await audit_service.log_event(
        db,
        action="chunk.alignment_fetched",
        resource_type="document_chunk",
        user_id=str(user_id),
        resource_id=str(chunk_id),
        details={"document_id": str(document_id), "voice_id": voice_id, "outcome": outcome},
        ip_address=request.client.host if request.client else None,
    )
    if not chunk_row or not chunk_row.text_content:
        raise HTTPException(status_code=404, detail="Chunk not found")

    source_type = (getattr(chunk_row, "source_type", "") or "").lower()
    chunk_text = chunk_row.text_content
    if source_type == "pdf":
        chunk_text = _clean_pdf_chunk_text(chunk_text, getattr(chunk_row, "page_number", None))

    # Cache invalidation on edit is already handled by
    # _invalidate_chunk_audio_cache, so the previous
    # `chunk_text == chunk_row.text_content` predicate was paranoid AND
    # broken (PDF cleaning made it always fail, see Bug-EL-1).
    cached = await get_alignment(str(chunk_id), voice_id)
    if cached is not None:
        logger.info(
            "audio.cache_hit",
            chunk_id=str(chunk_id),
            voice_id=voice_id,
            source_type=source_type,
            user_id=str(user_id),
            kind="alignment",
        )
        return cached

    # Synthesize with optional alignment, quota-aware (graceful Edge fallback).
    from psitta.providers.tts_router import TTSRouter

    tts = TTSRouter()
    try:
        audio_bytes, alignment, provider = await tts.synthesize_with_alignment_and_quota(
            chunk_text,
            voice_id,
            user_id=user_id,
            db=db,
        )
    except Exception as e:
        logger.error("audio.alignment_failed", error=str(e), voice_id=voice_id)
        raise HTTPException(status_code=502, detail=f"TTS alignment failed: {e}")

    # ElevenLabs and Edge both produce valid alignment for the audio they
    # returned (ElevenLabs via /with-timestamps; Edge via WordBoundary
    # chunks expanded server-side to char-level in the ElevenLabs schema).
    # Discard alignment from any other provider so a future Azure path
    # can't ship mismatched timestamps against its audio.
    if provider not in {"elevenlabs", "edge"} and alignment is not None:
        logger.warning(
            "audio.alignment_discarded",
            provider=provider,
            voice_id=voice_id,
            reason="alignment only valid for elevenlabs/edge",
        )
        alignment = None

    # Persist mp3 + alignment to local + S3
    await put_mp3(str(chunk_id), voice_id, audio_bytes)

    payload = {
        "document_id": str(document_id),
        "chunk_id": str(chunk_id),
        "voice_id": voice_id,
        "provider": provider,
        "alignment": alignment,
    }
    await put_alignment(str(chunk_id), voice_id, payload)

    # Ensure audio_segments cache record exists/updated (best-effort).
    # We keep the existing schema; duration is approximate if unknown.
    try:
        result = await db.execute(
            text("SELECT storage_key FROM audio_segments WHERE chunk_id = :cid AND voice_id = :vid"),
            {"cid": chunk_id, "vid": voice_id},
        )
        row = result.first()
        if row:
            await db.execute(
                text("DELETE FROM audio_segments WHERE chunk_id = :cid AND voice_id = :vid"),
                {"cid": chunk_id, "vid": voice_id},
            )
        from uuid import uuid4 as _uuid4

        await db.execute(
            text(
                "INSERT INTO audio_segments (id, document_id, chunk_id, voice_id, speed, storage_key, duration_ms, file_size_bytes, format) "
                "VALUES (:id, :doc_id, :chunk_id, :voice_id, :speed, :key, :dur, :size, :fmt)"
            ),
            {
                "id": _uuid4(),
                "doc_id": document_id,
                "chunk_id": chunk_id,
                "voice_id": voice_id,
                "speed": 1.0,
                "key": s3_key_mp3(str(chunk_id), voice_id),
                "dur": int(len(audio_bytes) / 24),
                "size": len(audio_bytes),
                "fmt": "mp3",
            },
        )
        await db.commit()
    except Exception as e:
        logger.warning("audio.alignment_cache_db_failed", error=str(e), chunk_id=str(chunk_id), voice_id=voice_id)

    return payload


async def _invalidate_chunk_audio_cache(chunk_id: UUID, db: AsyncSession) -> None:
    """Delete all audio_segments rows for a chunk so next request re-synthesizes.

    Clears all three cache layers:
      1. S3/MinIO objects (mp3 + alignment sidecar)
      2. Local /tmp cache files
      3. Database audio_segments rows
    """
    import glob as _glob
    import os
    from psitta.services.audio_cache import AUDIO_DIR

    logger.info("cache_invalidation.start", chunk_id=str(chunk_id))

    # 1. Delete ALL S3 objects with prefix audio/{chunk_id}_
    #    This doesn't depend on audio_segments DB rows existing.
    try:
        from psitta.config import get_settings
        from psitta.providers.storage_s3 import S3StorageProvider
        s3 = S3StorageProvider(get_settings())
        bucket = get_settings().S3_BUCKET_NAME
        prefix = f"audio/{chunk_id}_"
        deleted_count = await s3.delete_by_prefix(bucket, prefix)
        logger.info("cache_invalidation.s3_done", chunk_id=str(chunk_id), prefix=prefix, deleted=deleted_count)
    except Exception as e:
        logger.error("cache_invalidation.s3_failed", chunk_id=str(chunk_id), error=str(e))

    # 2. Delete local cache files
    pattern = os.path.join(AUDIO_DIR, f"{chunk_id}_*")
    matched = _glob.glob(pattern)
    logger.info("cache_invalidation.local_files", pattern=pattern, matched=matched)
    for path in matched:
        try:
            os.remove(path)
        except OSError as e:
            logger.error("cache_invalidation.local_delete_failed", path=path, error=str(e))

    # 3. Delete DB rows
    result = await db.execute(
        text("DELETE FROM audio_segments WHERE chunk_id = :cid"),
        {"cid": chunk_id},
    )
    logger.info("cache_invalidation.db_delete", chunk_id=str(chunk_id), rows_deleted=result.rowcount)


def _summarize_formatted_content(
    formatted_content: list[dict] | None,
) -> dict:
    """Privacy-respecting structural summary of formatted_content for
    ops logging. Captures block types, heading levels, list types,
    alignment, runs counts, and per-block flags signaling whether
    color/strike/font_family attributes are present on any run.

    NEVER includes run text content, color hex values, or font family
    names — flags only on M13.4 attributes. Truncates the blocks list
    at 50 entries with a marker so large documents don't bloat log
    lines.

    Returns:
      {
        "total_blocks": int,
        "blocks": [
          {"type": "heading", "level": int|None,
           "level_runtime_type": "int"|"str"|"NoneType"|...,
           "alignment": str|None,
           "runs_count": int,
           "has_color": bool, "has_strike": bool,
           "has_font_family": bool},
          {"type": "list_item", "list_type": "bullet"|"numbered"|None,
           "alignment": str|None,
           "runs_count": int,
           "has_color": bool, "has_strike": bool,
           "has_font_family": bool},
          {"type": "paragraph",
           "alignment": str|None,
           "runs_count": int,
           "has_color": bool, "has_strike": bool,
           "has_font_family": bool},
        ],
        "truncated_at": 50  # only present when total_blocks > 50
      }

    Defensive: returns {"total_blocks": 0, "blocks": []} on None input
    or any non-list structure.
    """
    if not isinstance(formatted_content, list):
        return {"total_blocks": 0, "blocks": []}

    summary_blocks: list[dict] = []
    total = len(formatted_content)
    for block in formatted_content[:50]:
        if not isinstance(block, dict):
            continue
        btype = block.get("type", "paragraph")
        runs = block.get("runs")
        runs_count = len(runs) if isinstance(runs, list) else 0

        # M13.4 Ship 1: per-block run-level attribute flags. We log
        # PRESENCE only — never values — so color hex codes and font
        # family names stay out of CloudWatch. Defensive isinstance
        # guards prevent a malformed block from crashing the
        # summarizer (which would lose forensic value of the log line).
        has_color = False
        has_strike = False
        has_font_family = False
        if isinstance(runs, list):
            for r in runs:
                if not isinstance(r, dict):
                    continue
                if r.get("strike") is True:
                    has_strike = True
                if isinstance(r.get("color"), str) and r.get("color"):
                    has_color = True
                if (
                    isinstance(r.get("font_family"), str)
                    and r.get("font_family")
                ):
                    has_font_family = True

        entry: dict = {
            "type": btype,
            "runs_count": runs_count,
            "alignment": block.get("alignment")
            if isinstance(block.get("alignment"), str)
            else None,
            "has_color": has_color,
            "has_strike": has_strike,
            "has_font_family": has_font_family,
        }
        if btype == "heading":
            raw_level = block.get("level")
            entry["level"] = raw_level if isinstance(raw_level, int) else None
            entry["level_runtime_type"] = type(raw_level).__name__
        elif btype == "list_item":
            entry["list_type"] = block.get("list_type")
        summary_blocks.append(entry)

    result: dict = {"total_blocks": total, "blocks": summary_blocks}
    if total > 50:
        result["truncated_at"] = 50
    return result


def _rebuild_formatted_content_for_chunk(
    text_content: str,
    existing_formatted: list[dict] | None,
) -> list[dict]:
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text_content) if p.strip()]
    if not paragraphs:
        return []

    existing_blocks = existing_formatted if isinstance(existing_formatted, list) else []
    rebuilt: list[dict] = []

    for index, paragraph in enumerate(paragraphs):
        previous = existing_blocks[index] if index < len(existing_blocks) and isinstance(existing_blocks[index], dict) else {}
        previous_runs = previous.get("runs") if isinstance(previous.get("runs"), list) else []
        first_run = previous_runs[0] if previous_runs and isinstance(previous_runs[0], dict) else {}

        block_type = previous.get("type", "paragraph")
        block: dict = {
            "type": block_type,
            "runs": [{
                "text": paragraph,
                "bold": bool(first_run.get("bold", False)),
                "italic": bool(first_run.get("italic", False)),
                "underline": bool(first_run.get("underline", False)),
                "font_size": first_run.get("font_size"),
            }],
        }
        if previous.get("level") is not None:
            block["level"] = previous["level"]
        rebuilt.append(block)

    return rebuilt


@router.patch("/{document_id}/chunks/{chunk_id}", response_model=ChunkResponse)
async def update_chunk_text(
    document_id: UUID,
    chunk_id: UUID,
    request: ChunkUpdateRequest,
    http_request: Request,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> ChunkResponse:
    """Update the text content of a chunk. Stores original text on first edit."""
    result = await db.execute(
        text(
            "SELECT c.id, c.sequence_index, c.chunk_type, c.text_content, c.tone, c.page_number, "
            "c.character_count, c.is_edited, c.edited_at, c.original_text, c.metadata_json, c.formatted_content "
            "FROM document_chunks c "
            "JOIN documents d ON d.id = c.document_id "
            "WHERE c.id = :cid AND c.document_id = :did "
            "AND d.user_id = :uid AND d.status != 'deleted'"
        ),
        {"cid": chunk_id, "did": document_id, "uid": str(user_id)},
    )
    chunk = result.first()
    outcome = "found" if chunk else "not_found_or_unauthorized"
    await audit_service.log_event(
        db,
        action="chunk.text_updated",
        resource_type="document_chunk",
        user_id=str(user_id),
        resource_id=str(chunk_id),
        details={"document_id": str(document_id), "outcome": outcome},
        ip_address=http_request.client.host if http_request.client else None,
    )
    if not chunk:
        raise HTTPException(status_code=404, detail="Chunk not found")

    import re as _re
    import unicodedata

    # Normalize unicode, remove invisible characters that break TTS
    new_text = request.text.strip()
    new_text = new_text.replace('\u200B', '').replace('\u200C', '')
    new_text = new_text.replace('\u200D', '').replace('\u00AD', '')
    new_text = new_text.replace('\uFEFF', '').replace('\u00A0', ' ')
    new_text = _re.sub(r' {2,}', ' ', new_text)
    new_text = unicodedata.normalize('NFC', new_text)

    now = datetime.now(timezone.utc)
    existing_meta = chunk.metadata_json if isinstance(chunk.metadata_json, dict) else {}
    segmenter = pysbd.Segmenter(language="en", clean=False)
    sentences = segmenter.segment(new_text)
    boundaries: list[list[int]] = []
    cursor = 0
    for sentence in sentences:
        start = new_text.index(sentence, cursor)
        end = start + len(sentence)
        boundaries.append([start, end])
        cursor = end

    updated_meta = {**existing_meta, "sentence_boundaries": boundaries}

    # Log a sanitized structural summary of the incoming formatted_content
    # for ops debugging — block types, heading levels, list_types, and runs
    # counts only. NEVER includes run text content. Wrapped in try/except so
    # a malformed payload can never break the API request.
    try:
        _fc_summary = _summarize_formatted_content(request.formatted_content)
        logger.info(
            "chunk.update.formatted_content_received",
            document_id=str(document_id),
            chunk_id=str(chunk_id),
            summary=_fc_summary,
        )
    except Exception as _fc_log_err:
        logger.warning(
            "chunk.update.summary_failed",
            document_id=str(document_id),
            chunk_id=str(chunk_id),
            error=str(_fc_log_err),
        )

    # Prefer client-authored formatted_content when supplied (the Phase 1
    # toolbar-persist path). Falls back to the legacy server-side rebuild
    # that inherits first-run attributes from pre-edit state when the
    # caller sent only plain text, preserving backward compatibility.
    if request.formatted_content is not None:
        updated_formatted = request.formatted_content
    else:
        updated_formatted = _rebuild_formatted_content_for_chunk(
            new_text,
            chunk.formatted_content if isinstance(chunk.formatted_content, list) else None,
        )
    meta_json = json.dumps(updated_meta, ensure_ascii=False)
    fmt_json = json.dumps(updated_formatted, ensure_ascii=False) if updated_formatted else None

    # Store original text on first edit only
    if not chunk.is_edited:
        await db.execute(
            text(
                "UPDATE document_chunks "
                "SET text_content = :txt, character_count = :chars, "
                "metadata_json = CAST(:meta AS jsonb), "
                "formatted_content = CAST(:fmt AS jsonb), "
                "is_edited = true, edited_at = :now, original_text = :orig "
                "WHERE id = :cid"
            ),
            {
                "cid": chunk_id,
                "txt": new_text,
                "chars": len(new_text),
                "meta": meta_json,
                "fmt": fmt_json,
                "now": now,
                "orig": chunk.text_content,
            },
        )
    else:
        await db.execute(
            text(
                "UPDATE document_chunks "
                "SET text_content = :txt, character_count = :chars, "
                "metadata_json = CAST(:meta AS jsonb), "
                "formatted_content = CAST(:fmt AS jsonb), "
                "edited_at = :now "
                "WHERE id = :cid"
            ),
            {
                "cid": chunk_id,
                "txt": new_text,
                "chars": len(new_text),
                "meta": meta_json,
                "fmt": fmt_json,
                "now": now,
            },
        )

    # Invalidate audio cache for this chunk
    await _invalidate_chunk_audio_cache(chunk_id, db)

    await db.commit()

    logger.info("chunk.updated", chunk_id=str(chunk_id), document_id=str(document_id))

    await audit_service.log_event(
        db,
        action="document.chunk.update",
        resource_type="document_chunk",
        user_id=str(user_id),
        resource_id=str(chunk_id),
        details={"document_id": str(document_id), "character_count": len(new_text)},
        ip_address=http_request.client.host if http_request.client else None,
    )
    return ChunkResponse(
        id=str(chunk.id),
        sequence_index=chunk.sequence_index,
        chunk_type=chunk.chunk_type if isinstance(chunk.chunk_type, str) else str(chunk.chunk_type),
        text_content=new_text,
        tone=chunk.tone if isinstance(chunk.tone, str) else str(chunk.tone),
        page_number=chunk.page_number,
        character_count=len(new_text),
        is_edited=True,
        edited_at=now,
        original_text=chunk.original_text if chunk.is_edited else chunk.text_content,
    )


@router.post("/{document_id}/chunks", response_model=ChunkResponse, status_code=status.HTTP_201_CREATED)
async def create_chunk(
    document_id: UUID,
    request: ChunkCreateRequest,
    http_request: Request,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> ChunkResponse:
    """Insert a new chunk into an existing document.

    M13.1b save-path helper. The client supplies its desired
    sequence_index; if a chunk already occupies that slot, the handler
    shifts the tail up by 100 000 to avoid colliding with the UNIQUE
    (document_id, sequence_index) constraint. The final authoritative
    order is established by the subsequent PATCH /documents/{id} call
    carrying the full chunk_positions list — the caller is expected to
    make that call atomically after fanning out all inserts/deletes.
    """
    # Authorization — confirm the caller owns the document.
    doc_result = await db.execute(
        text(
            "SELECT id FROM documents "
            "WHERE id = :did AND user_id = :uid AND status != 'deleted'"
        ),
        {"did": document_id, "uid": str(user_id)},
    )
    if doc_result.first() is None:
        raise HTTPException(status_code=404, detail="Document not found")

    # Normalize text (same rules as update_chunk_text).
    import re as _re
    import unicodedata
    new_text = request.text.strip()
    new_text = new_text.replace('​', '').replace('‌', '')
    new_text = new_text.replace('‍', '').replace('­', '')
    new_text = new_text.replace('﻿', '').replace(' ', ' ')
    new_text = _re.sub(r' {2,}', ' ', new_text)
    new_text = unicodedata.normalize('NFC', new_text)

    # Sentence boundaries (matches update_chunk_text pipeline).
    segmenter = pysbd.Segmenter(language="en", clean=False)
    sentences = segmenter.segment(new_text)
    boundaries: list[list[int]] = []
    cursor = 0
    for sentence in sentences:
        start = new_text.index(sentence, cursor)
        end = start + len(sentence)
        boundaries.append([start, end])
        cursor = end

    meta_json = json.dumps({"sentence_boundaries": boundaries}, ensure_ascii=False)
    fmt_json = (
        json.dumps(request.formatted_content, ensure_ascii=False)
        if request.formatted_content is not None
        else None
    )

    # Shift the tail up by 100 000 if the requested sequence_index is
    # occupied. This cheap two-pass pattern avoids the UNIQUE collision
    # and leaves the document in a readable (if temporarily non-dense)
    # state that the final PATCH /documents/{id} reindex will compact.
    await db.execute(
        text(
            "UPDATE document_chunks SET sequence_index = sequence_index + 100000 "
            "WHERE document_id = :did AND sequence_index >= :seq"
        ),
        {"did": document_id, "seq": request.sequence_index},
    )

    new_id = uuid4()
    now = datetime.now(timezone.utc)
    await db.execute(
        text(
            "INSERT INTO document_chunks "
            "(id, document_id, sequence_index, text_content, character_count, "
            "metadata_json, formatted_content, is_edited, edited_at, original_text, created_at) "
            "VALUES "
            "(:id, :did, :seq, :txt, :chars, CAST(:meta AS jsonb), CAST(:fmt AS jsonb), "
            "true, :now, NULL, :now)"
        ),
        {
            "id": new_id,
            "did": document_id,
            "seq": request.sequence_index,
            "txt": new_text,
            "chars": len(new_text),
            "meta": meta_json,
            "fmt": fmt_json,
            "now": now,
        },
    )

    await db.commit()

    logger.info(
        "document.chunk.created",
        chunk_id=str(new_id),
        document_id=str(document_id),
        sequence_index=request.sequence_index,
    )

    await audit_service.log_event(
        db,
        action="document.chunk.create",
        resource_type="document_chunk",
        user_id=str(user_id),
        resource_id=str(new_id),
        details={
            "document_id": str(document_id),
            "sequence_index": request.sequence_index,
            "character_count": len(new_text),
        },
        ip_address=http_request.client.host if http_request.client else None,
    )

    return ChunkResponse(
        id=str(new_id),
        sequence_index=request.sequence_index,
        chunk_type="text",
        text_content=new_text,
        tone="neutral",
        page_number=request.page_number or 1,
        character_count=len(new_text),
        is_edited=True,
        edited_at=now,
        original_text=None,
    )


@router.delete("/{document_id}/chunks/{chunk_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_chunk(
    document_id: UUID,
    chunk_id: UUID,
    http_request: Request,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> None:
    """Delete a chunk from a document.

    M13.1b save-path helper. Invalidates the three-layer audio cache
    (S3 + local /tmp + audio_segments DB rows) BEFORE the chunk row
    is removed. The audio_segments FK has ON DELETE CASCADE so the DB
    layer would purge automatically; the explicit call ensures S3 and
    /tmp objects (which the FK cannot touch) are also cleared while
    the chunk_id is still a valid reference.
    """
    # Authorization — confirm the caller owns the document and the
    # chunk belongs to it. One JOIN query rules out both 404 cases.
    check = await db.execute(
        text(
            "SELECT c.id FROM document_chunks c "
            "JOIN documents d ON d.id = c.document_id "
            "WHERE c.id = :cid AND d.id = :did "
            "AND d.user_id = :uid AND d.status != 'deleted'"
        ),
        {"cid": chunk_id, "did": document_id, "uid": str(user_id)},
    )
    if check.first() is None:
        raise HTTPException(status_code=404, detail="Chunk not found")

    await _invalidate_chunk_audio_cache(chunk_id, db)

    await db.execute(
        text(
            "DELETE FROM document_chunks "
            "WHERE id = :cid AND document_id = :did"
        ),
        {"cid": chunk_id, "did": document_id},
    )
    await db.commit()

    logger.info(
        "document.chunk.deleted",
        chunk_id=str(chunk_id),
        document_id=str(document_id),
    )
    await audit_service.log_event(
        db,
        action="document.chunk.delete",
        resource_type="document_chunk",
        user_id=str(user_id),
        resource_id=str(chunk_id),
        details={"document_id": str(document_id)},
        ip_address=http_request.client.host if http_request.client else None,
    )


@router.post("/{document_id}/chunks/{chunk_id}/resynthesize", response_model=ResynthesizeResponse)
async def resynthesize_chunk(
    document_id: UUID,
    chunk_id: UUID,
    request: Request,
    voice_id: str = Query(default="21m00Tcm4TlvDq8ikWAM"),
    speed: float = Query(default=1.0),
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> ResynthesizeResponse:
    """Re-synthesize audio for an edited chunk using its current text_content."""
    result = await db.execute(
        text(
            "SELECT c.id FROM document_chunks c "
            "JOIN documents d ON d.id = c.document_id "
            "WHERE c.id = :cid AND c.document_id = :did "
            "AND d.user_id = :uid AND d.status != 'deleted'"
        ),
        {"cid": chunk_id, "did": document_id, "uid": str(user_id)},
    )
    chunk_owned = result.first() is not None
    outcome = "found" if chunk_owned else "not_found_or_unauthorized"
    await audit_service.log_event(
        db,
        action="chunk.resynthesized",
        resource_type="document_chunk",
        user_id=str(user_id),
        resource_id=str(chunk_id),
        details={"document_id": str(document_id), "voice_id": voice_id, "outcome": outcome},
        ip_address=request.client.host if request.client else None,
    )
    if not chunk_owned:
        raise HTTPException(status_code=404, detail="Chunk not found")

    # Invalidate cache so next audio request re-synthesizes
    await _invalidate_chunk_audio_cache(chunk_id, db)
    await db.commit()

    audio_url = f"/api/v1/documents/{document_id}/chunks/{chunk_id}/audio?voice_id={voice_id}&speed={speed}"

    logger.info("chunk.resynthesize", chunk_id=str(chunk_id), voice_id=voice_id)

    await audit_service.log_event(
        db,
        action="document.chunk.resynthesize",
        resource_type="document_chunk",
        user_id=str(user_id),
        resource_id=str(chunk_id),
        details={"document_id": str(document_id), "voice_id": voice_id, "speed": speed},
        ip_address=request.client.host if request.client else None,
    )
    return ResynthesizeResponse(
        chunk_id=str(chunk_id),
        audio_url=audio_url,
        message="Cache invalidated. Next audio request will re-synthesize with updated text.",
    )


@router.post("/{document_id}/resynthesize", status_code=202)
async def resynthesize_document(
    document_id: UUID,
    background_tasks: BackgroundTasks,
    request: Request,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> JSONResponse:
    """Clear audio cache for all chunks of a document and queue re-synthesis.

    Returns 202 Accepted immediately. All cache invalidation and
    re-synthesis run in a background task.
    """
    # Validate document exists and belongs to the authenticated user
    doc_result = await db.execute(
        text(
            "SELECT id FROM documents "
            "WHERE id = :did AND user_id = :uid AND status != 'deleted'"
        ),
        {"did": document_id, "uid": str(user_id)},
    )
    if not doc_result.first():
        raise HTTPException(status_code=404, detail="Document not found")

    # Fetch all chunk IDs for this document
    chunk_result = await db.execute(
        text(
            "SELECT id FROM document_chunks "
            "WHERE document_id = :did ORDER BY sequence_index"
        ),
        {"did": document_id},
    )
    chunk_ids = [row[0] for row in chunk_result.fetchall()]
    chunk_count = len(chunk_ids)

    logger.info(
        "document.resynthesize.queued",
        document_id=str(document_id),
        chunk_count=chunk_count,
    )

    # All cache invalidation + re-synthesis runs in background
    background_tasks.add_task(
        _background_invalidate_and_resynthesize, document_id, chunk_ids
    )

    await audit_service.log_event(
        db,
        action="document.resynthesize",
        resource_type="document",
        user_id=str(user_id),
        resource_id=str(document_id),
        details={"chunk_count": chunk_count},
        ip_address=request.client.host if request.client else None,
    )
    return JSONResponse(
        status_code=202,
        content={
            "document_id": str(document_id),
            "chunks_cleared": chunk_count,
            "message": "Audio regeneration started in background",
        },
    )


async def _background_invalidate_and_resynthesize(
    document_id: UUID, chunk_ids: list[UUID]
) -> None:
    """Background task: invalidate audio cache for all chunks, then re-synthesize."""
    try:
        from psitta.db.session import async_session_factory

        async with async_session_factory() as db:
            for cid in chunk_ids:
                await _invalidate_chunk_audio_cache(cid, db)
            await db.commit()

        logger.info(
            "document.resynthesize.cache_cleared",
            document_id=str(document_id),
            chunks_cleared=len(chunk_ids),
        )

        if chunk_ids:
            await _eager_synthesize_chunks(document_id, chunk_ids)

        logger.info(
            "document.resynthesize.complete",
            document_id=str(document_id),
        )
    except Exception as e:
        logger.error(
            "document.resynthesize.failed",
            document_id=str(document_id),
            error=str(e),
        )


# ── POST /{document_id}/summarize ─────────────────────────────────────────────


@router.post("/{document_id}/summarize")
async def summarize_document(
    document_id: UUID,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> dict:
    """Summarize a document using LLM (Writing Nook Pro / Creative Nook Pro).

    Performs a quota pre-check against llm_usage_counters, calls the
    OpenAI provider, increments the counter, and returns the summary
    with token accounting for the current billing period.

    Raises 402 if the plan has no LLM access or the period quota is exhausted.
    """
    from psitta.services.llm_service import summarize_with_quota

    return await summarize_with_quota(db, user_id, document_id)


class DocumentUpdateRequest(BaseModel):
    title: str | None = Field(None, min_length=1, max_length=200)
    cover_type: str | None = None
    cover_value: str | None = None
    # M13.1b: authoritative chunk-offset map. When present, the handler
    # persists it verbatim to documents.chunk_positions AND rebuilds
    # document_chunks.sequence_index from the supplied ordering in the
    # same transaction. List shape: [{"chunk_id": str, "start_offset":
    # int, "end_offset": int}, ...].
    chunk_positions: list[dict] | None = None
    chunk_count: int | None = Field(default=None, ge=0)


@router.patch("/{document_id}")
async def update_document(
    document_id: UUID,
    payload: DocumentUpdateRequest,
    request: Request,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> dict:
    """Update editable document fields (title, cover_type, cover_value)."""
    # Build dynamic SET clause based on provided fields
    set_parts: list[str] = []
    params: dict = {"did": document_id, "uid": str(user_id)}
    updated_fields: list[str] = []

    if payload.title is not None:
        set_parts.append("title = :title")
        params["title"] = payload.title.strip()
        updated_fields.append("title")

    # Use model_fields_set to detect explicitly-sent fields (including null for "remove")
    cover_type_sent = 'cover_type' in payload.model_fields_set
    cover_value_sent = 'cover_value' in payload.model_fields_set

    if cover_type_sent or cover_value_sent:
        # Fetch current cover info to check if we need to delete old uploaded cover
        cur_result = await db.execute(
            text("SELECT cover_type, cover_value FROM documents WHERE id = :did AND user_id = :uid AND status != 'deleted'"),
            {"did": document_id, "uid": str(user_id)},
        )
        cur_row = cur_result.first()
        if not cur_row:
            raise HTTPException(status_code=404, detail="Document not found")

        old_cover_type = cur_row.cover_type
        old_cover_value = cur_row.cover_value

        if cover_type_sent:
            set_parts.append("cover_type = :cover_type")
            params["cover_type"] = payload.cover_type
            updated_fields.append("cover_type")

        if cover_value_sent:
            set_parts.append("cover_value = :cover_value")
            params["cover_value"] = payload.cover_value
            updated_fields.append("cover_value")

        # If old cover was uploaded and we're changing away from it, delete from S3
        new_cover_type = payload.cover_type if cover_type_sent else old_cover_type
        if old_cover_type == "uploaded" and new_cover_type != "uploaded" and old_cover_value:
            try:
                from psitta.config import get_settings
                from psitta.providers.storage_s3 import S3StorageProvider
                settings = get_settings()
                s3 = S3StorageProvider(settings)
                bucket = settings.S3_BUCKET_NAME
                await s3.delete_by_prefix(bucket, f"covers/{document_id}.")
                logger.info("document.cover.s3_deleted", doc_id=str(document_id))
            except Exception as e:
                logger.error("document.cover.s3_delete_failed", doc_id=str(document_id), error=str(e))

    # M13.1b — chunk_positions + chunk_count extension ──────────────────
    chunk_positions_sent = 'chunk_positions' in payload.model_fields_set
    chunk_count_sent = 'chunk_count' in payload.model_fields_set
    chunk_positions = payload.chunk_positions if chunk_positions_sent else None

    if chunk_positions_sent and chunk_positions is not None:
        # Validate shape: each entry must have chunk_id, start_offset,
        # end_offset; reject malformed input before doing any writes so
        # the document never ends up with a partial reindex.
        for entry in chunk_positions:
            if not isinstance(entry, dict):
                raise HTTPException(status_code=422, detail="chunk_positions entries must be objects")
            if 'chunk_id' not in entry or 'start_offset' not in entry or 'end_offset' not in entry:
                raise HTTPException(
                    status_code=422,
                    detail="each chunk_positions entry requires chunk_id, start_offset, end_offset",
                )
        # Cross-check count when supplied — guards against client drift.
        if chunk_count_sent and payload.chunk_count is not None and len(chunk_positions) != payload.chunk_count:
            raise HTTPException(
                status_code=422,
                detail=f"chunk_positions length ({len(chunk_positions)}) does not match chunk_count ({payload.chunk_count})",
            )
        set_parts.append("chunk_positions = CAST(:chunk_positions AS jsonb)")
        params["chunk_positions"] = json.dumps(chunk_positions, ensure_ascii=False)
        updated_fields.append("chunk_positions")

    if chunk_count_sent and payload.chunk_count is not None:
        # chunk_count isn't a real column on documents today; the
        # canonical source is the COUNT(*) over document_chunks.
        # Accept the field so the client can assert consistency; track
        # it in updated_fields for audit but skip the SET clause.
        updated_fields.append("chunk_count")

    if not set_parts:
        raise HTTPException(status_code=400, detail="No fields to update")

    set_parts.append("updated_at = NOW()")
    set_clause = ", ".join(set_parts)

    result = await db.execute(
        text(
            f"UPDATE documents SET {set_clause} "
            "WHERE id = :did AND user_id = :uid AND status != 'deleted'"
        ),
        params,
    )
    if result.rowcount == 0:
        raise HTTPException(status_code=404, detail="Document not found")

    # M13.1b — reindex document_chunks.sequence_index in the same
    # transaction. Two-pass pattern to avoid the UNIQUE (document_id,
    # sequence_index) collision: first move everyone out of the
    # destination range, then assign final indices in the supplied
    # chunk_positions order.
    if chunk_positions_sent and chunk_positions is not None and len(chunk_positions) > 0:
        await db.execute(
            text(
                "UPDATE document_chunks SET sequence_index = sequence_index + 100000 "
                "WHERE document_id = :did"
            ),
            {"did": document_id},
        )
        for i, entry in enumerate(chunk_positions):
            await db.execute(
                text(
                    "UPDATE document_chunks SET sequence_index = :idx "
                    "WHERE id = :cid AND document_id = :did"
                ),
                {"idx": i, "cid": entry['chunk_id'], "did": document_id},
            )

    await db.commit()

    row = await db.execute(
        text(
            "SELECT id, user_id, title, status, source_type, page_count, word_count, "
            "file_size_bytes, created_at, updated_at, cover_type, cover_value "
            "FROM documents WHERE id = :did"
        ),
        {"did": document_id},
    )
    doc = row.first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    logger.info("document.updated", doc_id=str(document_id), fields=updated_fields)

    await audit_service.log_event(
        db,
        action="document.update",
        resource_type="document",
        user_id=str(user_id),
        resource_id=str(document_id),
        details={"fields": updated_fields},
        ip_address=request.client.host if request.client else None,
    )
    return {
        "id": str(doc.id),
        "title": doc.title,
        "status": doc.status,
        "source_type": doc.source_type,
        "page_count": doc.page_count,
        "word_count": getattr(doc, "word_count", 0),
        "file_size_bytes": doc.file_size_bytes,
        "created_at": doc.created_at.isoformat() if doc.created_at else None,
        "updated_at": doc.updated_at.isoformat() if doc.updated_at else None,
        "cover_type": doc.cover_type,
        "cover_value": doc.cover_value,
    }


@router.post("/{document_id}/cover")
async def upload_cover(
    document_id: UUID,
    file: UploadFile,
    request: Request,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> dict:
    """Upload a cover image for a document."""
    # Validate document exists
    doc_result = await db.execute(
        text(
            "SELECT id, cover_type, cover_value FROM documents "
            "WHERE id = :did AND user_id = :uid AND status != 'deleted'"
        ),
        {"did": document_id, "uid": str(user_id)},
    )
    doc_row = doc_result.first()
    if not doc_row:
        raise HTTPException(status_code=404, detail="Document not found")

    # Validate file type by extension
    filename = file.filename or "unknown"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    allowed_image_exts = {"jpeg", "jpg", "gif", "png"}
    if ext not in allowed_image_exts:
        raise HTTPException(status_code=415, detail="Unsupported image type. Allowed: jpeg, jpg, gif, png")

    # Validate content type
    content_type = file.content_type or ""
    allowed_content_types = {"image/jpeg", "image/jpg", "image/gif", "image/png"}
    if content_type not in allowed_content_types:
        raise HTTPException(status_code=415, detail=f"Unsupported content type: {content_type}")

    # Read and validate size (max 20MB)
    file_bytes = await file.read()
    max_size = 20 * 1024 * 1024
    if len(file_bytes) > max_size:
        raise HTTPException(status_code=413, detail="Image too large. Please use an image under 20MB.")

    # Always resize with Pillow (max 1600x1600, maintain aspect ratio). 1600 is
    # large enough to stay crisp at every display size (cards, detail, player,
    # and future full-size book-cover views) without storing the raw upload.
    try:
        from PIL import Image

        img = Image.open(io.BytesIO(file_bytes))

        # GIF: extract first frame only, convert to RGBA then save as PNG
        if ext == "gif":
            img = img.convert("RGBA")
            ext = "png"

        # Fit within 1600x1600 box maintaining aspect ratio
        max_dim = 1600
        if img.width > max_dim or img.height > max_dim:
            img.thumbnail((max_dim, max_dim), Image.LANCZOS)

        buf = io.BytesIO()
        out_format = "PNG" if ext == "png" else "JPEG"
        save_kwargs = {"quality": 88} if out_format == "JPEG" else {}
        if out_format == "JPEG" and img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        img.save(buf, format=out_format, **save_kwargs)
        file_bytes = buf.getvalue()
        logger.info("document.cover.resized", doc_id=str(document_id),
                     width=img.width, height=img.height, format=out_format,
                     size=len(file_bytes))
    except ImportError:
        logger.warning("document.cover.pillow_unavailable", doc_id=str(document_id))
    except Exception as e:
        logger.warning("document.cover.resize_failed", doc_id=str(document_id), error=str(e))

    # Delete old cover from S3 if it existed
    old_cover_type = doc_row.cover_type
    old_cover_value = doc_row.cover_value
    if old_cover_type == "uploaded" and old_cover_value:
        try:
            from psitta.config import get_settings
            from psitta.providers.storage_s3 import S3StorageProvider
            settings = get_settings()
            s3 = S3StorageProvider(settings)
            bucket = settings.S3_BUCKET_NAME
            await s3.delete_by_prefix(bucket, f"covers/{document_id}.")
            logger.info("document.cover.old_deleted", doc_id=str(document_id))
        except Exception as e:
            logger.error("document.cover.old_delete_failed", doc_id=str(document_id), error=str(e))

    # Store in S3
    storage_key = f"covers/{document_id}.{ext}"
    content_type_map = {"jpeg": "image/jpeg", "jpg": "image/jpeg", "gif": "image/gif", "png": "image/png"}

    from psitta.config import get_settings
    from psitta.providers.storage_s3 import S3StorageProvider
    settings = get_settings()
    s3 = S3StorageProvider(settings)
    bucket = settings.S3_BUCKET_NAME
    await s3.put_object(bucket, storage_key, file_bytes, content_type=content_type_map.get(ext, "image/jpeg"))

    # Update document record
    await db.execute(
        text(
            "UPDATE documents SET cover_type = :ct, cover_value = :cv, updated_at = NOW() "
            "WHERE id = :did"
        ),
        {"ct": "uploaded", "cv": storage_key, "did": document_id},
    )
    await db.commit()

    logger.info("document.cover.uploaded", doc_id=str(document_id), key=storage_key, size=len(file_bytes))

    await audit_service.log_event(
        db,
        action="document.cover.upload",
        resource_type="document",
        user_id=str(user_id),
        resource_id=str(document_id),
        details={"storage_key": storage_key, "size_bytes": len(file_bytes)},
        ip_address=request.client.host if request.client else None,
    )
    # Return updated document dict
    row = await db.execute(
        text(
            "SELECT id, title, status, source_type, page_count, word_count, "
            "file_size_bytes, created_at, updated_at, cover_type, cover_value "
            "FROM documents WHERE id = :did"
        ),
        {"did": document_id},
    )
    doc = row.first()
    return {
        "id": str(doc.id),
        "title": doc.title,
        "status": doc.status,
        "source_type": doc.source_type,
        "page_count": doc.page_count,
        "word_count": getattr(doc, "word_count", 0),
        "file_size_bytes": doc.file_size_bytes,
        "created_at": doc.created_at.isoformat() if doc.created_at else None,
        "updated_at": doc.updated_at.isoformat() if doc.updated_at else None,
        "cover_type": doc.cover_type,
        "cover_value": doc.cover_value,
    }


@router.get("/{document_id}/cover")
async def get_cover(
    document_id: UUID,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
):
    """Serve the uploaded cover image for a document."""
    from fastapi.responses import Response

    result = await db.execute(
        text(
            "SELECT cover_type, cover_value FROM documents "
            "WHERE id = :did AND user_id = :uid AND status != 'deleted'"
        ),
        {"did": document_id, "uid": str(user_id)},
    )
    row = result.first()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")

    if row.cover_type != "uploaded" or not row.cover_value:
        raise HTTPException(status_code=404, detail="No uploaded cover image")

    # Fetch from S3
    from psitta.config import get_settings
    from psitta.providers.storage_s3 import S3StorageProvider
    settings = get_settings()
    s3 = S3StorageProvider(settings)
    bucket = settings.S3_BUCKET_NAME

    try:
        image_bytes = await s3.get_object(bucket, row.cover_value)
    except Exception as e:
        logger.error("document.cover.fetch_failed", doc_id=str(document_id), error=str(e))
        raise HTTPException(status_code=404, detail="Cover image not found in storage")

    # Determine content type from key extension
    ext = row.cover_value.rsplit(".", 1)[-1].lower() if "." in row.cover_value else "jpeg"
    content_type_map = {"jpeg": "image/jpeg", "jpg": "image/jpeg", "gif": "image/gif", "png": "image/png"}
    media_type = content_type_map.get(ext, "image/jpeg")

    return Response(
        content=image_bytes,
        media_type=media_type,
        headers={"Cache-Control": "max-age=3600"},
    )


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    document_id: UUID,
    request: Request,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> None:
    result = await db.execute(
        text(
            "UPDATE documents SET status = 'deleted', updated_at = NOW() "
            "WHERE id = :did AND user_id = :uid AND status != 'deleted'"
        ),
        {"did": document_id, "uid": str(user_id)},
    )
    if result.rowcount == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    logger.info("document.deleted", doc_id=str(document_id))
    await audit_service.log_event(
        db,
        action="document.delete",
        resource_type="document",
        user_id=str(user_id),
        resource_id=str(document_id),
        details=None,
        ip_address=request.client.host if request.client else None,
    )


@router.post("/{document_id}/restore", status_code=status.HTTP_200_OK)
async def restore_document(
    document_id: UUID,
    request: Request,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> dict:
    """Restore a soft-deleted document from Trash back to the Library."""
    result = await db.execute(
        text(
            "UPDATE documents SET status = 'ready', updated_at = NOW() "
            "WHERE id = :did AND user_id = :uid AND status = 'deleted'"
        ),
        {"did": document_id, "uid": str(user_id)},
    )
    if result.rowcount == 0:
        raise HTTPException(status_code=404, detail="Document not found in Trash")
    await db.commit()
    logger.info("document.restored", doc_id=str(document_id))
    await audit_service.log_event(
        db,
        action="document.restore",
        resource_type="document",
        user_id=str(user_id),
        resource_id=str(document_id),
        details=None,
        ip_address=request.client.host if request.client else None,
    )
    return {"id": str(document_id), "status": "ready"}


@router.delete(
    "/{document_id}/permanent", status_code=status.HTTP_204_NO_CONTENT
)
async def purge_document(
    document_id: UUID,
    request: Request,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> None:
    """Permanently delete a document that is already in Trash.

    Hard-deletes the documents row; chunks, audio_segments and sessions cascade
    via their ON DELETE CASCADE foreign keys. S3 cleanup (raw file + cover) is
    best-effort — orphaned objects are harmless and must never block the purge.
    """
    row = (
        await db.execute(
            text(
                "SELECT storage_key, cover_type, cover_value FROM documents "
                "WHERE id = :did AND user_id = :uid AND status = 'deleted'"
            ),
            {"did": document_id, "uid": str(user_id)},
        )
    ).first()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found in Trash")

    try:
        from psitta.config import get_settings
        from psitta.providers.storage_s3 import S3StorageProvider

        settings = get_settings()
        s3 = S3StorageProvider(settings)
        bucket = settings.S3_BUCKET_NAME
        if row.storage_key:
            await s3.delete_by_prefix(bucket, row.storage_key)
        if row.cover_type == "uploaded":
            await s3.delete_by_prefix(bucket, f"covers/{document_id}.")
    except Exception:
        logger.warning(
            "document.purge.s3_cleanup_failed", doc_id=str(document_id)
        )

    await db.execute(
        text("DELETE FROM documents WHERE id = :did AND user_id = :uid"),
        {"did": document_id, "uid": str(user_id)},
    )
    await db.commit()
    logger.info("document.purged", doc_id=str(document_id))
    await audit_service.log_event(
        db,
        action="document.purge",
        resource_type="document",
        user_id=str(user_id),
        resource_id=str(document_id),
        details=None,
        ip_address=request.client.host if request.client else None,
    )


@router.patch("/{document_id}/archive", status_code=status.HTTP_200_OK)
async def archive_document(
    document_id: UUID,
    request: Request,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> dict:
    """Toggle document between archived and ready status."""
    result = await db.execute(
        text(
            "SELECT status FROM documents "
            "WHERE id = :did AND user_id = :uid AND status != 'deleted'"
        ),
        {"did": document_id, "uid": str(user_id)},
    )
    row = result.first()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    new_status = "ready" if row.status == "archived" else "archived"
    await db.execute(
        text(
            "UPDATE documents SET status = :st, updated_at = NOW() "
            "WHERE id = :did AND user_id = :uid"
        ),
        {"st": new_status, "did": document_id, "uid": str(user_id)},
    )
    await db.commit()
    logger.info("document.archived", doc_id=str(document_id), new_status=new_status)
    await audit_service.log_event(
        db,
        action="document.archive",
        resource_type="document",
        user_id=str(user_id),
        resource_id=str(document_id),
        details={"new_status": new_status},
        ip_address=request.client.host if request.client else None,
    )
    return {"id": str(document_id), "status": new_status}


@router.patch("/{document_id}/project")
async def assign_project(
    document_id: str,
    body: dict,
    request: Request,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
):
    """Assign or remove a document from a project.
    Body: {"project_id": "<uuid>"} to assign, {"project_id": null} to remove.
    """
    project_id = body.get("project_id")
    # Verify document exists and belongs to authenticated user
    row = await db.execute(
        text("SELECT id FROM documents WHERE id = :id AND user_id = :uid AND status != 'deleted'"),
        {"id": document_id, "uid": str(user_id)},
    )
    if not row.mappings().first():
        raise HTTPException(status_code=404, detail="Document not found")
    # Verify project exists if assigning
    if project_id is not None:
        proj_row = await db.execute(
            text("SELECT id FROM projects WHERE id = :id AND user_id = :uid"),
            {"id": project_id, "uid": str(user_id)},
        )
        if not proj_row.mappings().first():
            raise HTTPException(status_code=404, detail="Project not found")
    await db.execute(
        text("UPDATE documents SET project_id = :pid WHERE id = :id"),
        {"pid": project_id, "id": document_id},
    )
    await db.commit()
    await audit_service.log_event(
        db,
        action="document.assign_project",
        resource_type="document",
        user_id=str(user_id),
        resource_id=document_id,
        details={"project_id": project_id},
        ip_address=request.client.host if request.client else None,
    )
    return {"id": document_id, "project_id": project_id}


@router.get("/{document_id}/download")
async def download_document(
    document_id: UUID,
    request: Request,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> FileResponse:
    """Serve the original uploaded file for download."""
    result = await db.execute(
        text(
            "SELECT title, source_type, storage_key FROM documents "
            "WHERE id = :did AND user_id = :uid AND status != 'deleted'"
        ),
        {"did": document_id, "uid": str(user_id)},
    )
    row = result.first()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")

    from psitta.services.audio_cache import get_raw_file

    if not row.storage_key:
        raise HTTPException(status_code=404, detail="Original file not available for download")

    file_path = await get_raw_file(row.storage_key)
    if not file_path:
        raise HTTPException(status_code=404, detail="Original file not available for download")

    filename = f"{row.title}.{row.source_type}"
    await audit_service.log_event(
        db,
        action="document.download",
        resource_type="document",
        user_id=str(user_id),
        resource_id=str(document_id),
        details={"filename": filename, "source_type": row.source_type},
        ip_address=request.client.host if request.client else None,
    )
    return FileResponse(
        file_path,
        media_type="application/octet-stream",
        filename=filename,
    )


@router.post("/{document_id}/duplicate", status_code=201)
async def duplicate_document(
    document_id: UUID,
    request: Request,
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
) -> dict:
    """Copy a document (row + chunks + raw file + cover) into the Library.

    The copy keeps the same project and cover; the title gets a ' (Copy)'
    suffix. JSONB columns (metadata_json, chunk_positions, formatted_content)
    are copied natively via INSERT ... SELECT to avoid round-tripping JSON
    through Python. Duplicating does not consume the monthly upload quota —
    it is the writer's own content.
    """
    src = (
        await db.execute(
            text(
                "SELECT title, source_type, storage_key FROM documents "
                "WHERE id = :did AND user_id = :uid AND status != 'deleted'"
            ),
            {"did": document_id, "uid": str(user_id)},
        )
    ).first()
    if not src:
        raise HTTPException(status_code=404, detail="Document not found")

    new_id = uuid4()
    extension = f".{src.source_type}"
    now = datetime.now(timezone.utc)
    new_title = f"{src.title} (Copy)"
    new_storage_key = (
        f"uploads/{new_id}{extension}" if src.storage_key else None
    )

    from psitta.config import get_settings
    from psitta.providers.storage_s3 import S3StorageProvider

    settings = get_settings()
    s3 = S3StorageProvider(settings)
    bucket = settings.S3_BUCKET_NAME

    # 1. Copy the raw uploaded file (best-effort — chunks drive reading, so a
    #    missing raw file doesn't block the copy; it just disables download).
    if src.storage_key and new_storage_key:
        try:
            raw = await s3.get_object(bucket, src.storage_key)
            await s3.put_object(
                bucket, new_storage_key, raw, "application/octet-stream"
            )
        except Exception:
            logger.warning(
                "document.duplicate.raw_copy_failed", doc_id=str(document_id)
            )
            new_storage_key = None

    # 2. Copy the documents row natively; override only what changes. The copy
    #    starts with no cover (cover_type/value forced NULL) — the writer picks
    #    one in the Library.
    await db.execute(
        text(
            "INSERT INTO documents "
            "(id, user_id, title, source_type, status, file_size_bytes, "
            "storage_key, page_count, word_count, metadata_json, "
            "chunk_positions, project_id, cover_type, cover_value, "
            "created_at, updated_at) "
            "SELECT :new_id, user_id, :new_title, source_type, status, "
            "file_size_bytes, :new_key, page_count, word_count, metadata_json, "
            "chunk_positions, project_id, NULL, NULL, :now, :now "
            "FROM documents WHERE id = :src_id AND user_id = :uid"
        ),
        {
            "new_id": new_id,
            "new_title": new_title,
            "new_key": new_storage_key,
            "now": now,
            "src_id": document_id,
            "uid": str(user_id),
        },
    )

    # 3. Copy all chunks natively with fresh ids.
    await db.execute(
        text(
            "INSERT INTO document_chunks "
            "(id, document_id, sequence_index, chunk_type, text_content, "
            "tone, page_number, character_count, metadata_json, "
            "formatted_content) "
            "SELECT gen_random_uuid(), :new_id, sequence_index, chunk_type, "
            "text_content, tone, page_number, character_count, metadata_json, "
            "formatted_content "
            "FROM document_chunks WHERE document_id = :src_id"
        ),
        {"new_id": new_id, "src_id": document_id},
    )

    await audit_service.log_event(
        db,
        action="document.duplicate",
        resource_type="document",
        user_id=str(user_id),
        resource_id=str(new_id),
        details={"source_document_id": str(document_id), "title": new_title},
        ip_address=request.client.host if request.client else None,
    )
    logger.info(
        "document.duplicated",
        src_doc_id=str(document_id),
        new_doc_id=str(new_id),
    )
    return {
        "id": str(new_id),
        "title": new_title,
        "source_type": src.source_type,
        "cover_type": None,
        "cover_value": None,
    }


# ── Branded DOCX export ──────────────────────────────────────────────────────

_LOGO_PATH = Path(__file__).resolve().parents[1] / "assets" / "logo.png"


def _build_branded_docx(
    *,
    title: str,
    chunks: list[dict],
    project_name: str | None,
    include_cover: bool,
    include_footer: bool,
) -> bytes:
    """Build a branded DOCX from document chunks and return raw bytes."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = DocxDocument()

    # ── Style defaults ────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"
    style.paragraph_format.space_after = Pt(6)

    has_logo = _LOGO_PATH.exists()

    # ── Cover page ────────────────────────────────────────────────────
    if include_cover:
        # Spacer at top
        for _ in range(6):
            doc.add_paragraph()

        # Logo
        if has_logo:
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_para.add_run()
            logo_run.add_picture(str(_LOGO_PATH), width=Inches(2.0))

            doc.add_paragraph()

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(26)
        title_run.font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)

        # Project / author
        if project_name:
            proj_para = doc.add_paragraph()
            proj_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            proj_run = proj_para.add_run(project_name)
            proj_run.font.size = Pt(14)
            proj_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.now(timezone.utc).strftime("%B %d, %Y"))
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Bottom branding
        for _ in range(4):
            doc.add_paragraph()
        brand_para = doc.add_paragraph()
        brand_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        brand_run = brand_para.add_run("Processed by Psitta")
        brand_run.font.size = Pt(10)
        brand_run.font.italic = True
        brand_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Page break after cover
        doc.add_page_break()

    # ── Document content ──────────────────────────────────────────────
    for chunk in chunks:
        chunk_title = chunk.get("title", "")
        chunk_text = chunk.get("text_content", "")
        chunk_formatted = chunk.get("formatted_content")

        if chunk_title:
            heading = doc.add_heading(chunk_title, level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)

        # Prefer formatted_content (M13.1b) when present so Bold/Italic/
        # Underline/font_size and heading/list block types round-trip into
        # the downloaded DOCX. Fall back to text_content for pre-M13.1b
        # chunks where formatted_content IS NULL.
        if chunk_formatted:
            # M13.4 Ship 1: alignment maps to python-docx WD_ALIGN_PARAGRAPH;
            # only emit when the saved value is one of the four canonical
            # values so unknown shapes fall through to Word's style
            # cascade default rather than raising.
            _ALIGN_MAP = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            for block in chunk_formatted:
                if not isinstance(block, dict):
                    continue
                btype = block.get("type", "paragraph")
                if btype == "heading":
                    level = block.get("level", 2)
                    try:
                        level_int = max(1, min(int(level), 6))
                    except (ValueError, TypeError):
                        level_int = 2
                    para = doc.add_heading("", level=level_int)
                elif btype == "list_item":
                    list_style = "List Bullet" if block.get("list_type") == "bullet" else "List Number"
                    para = doc.add_paragraph("", style=list_style)
                else:
                    para = doc.add_paragraph("")

                # Block-level alignment (M13.4 Ship 1) — composes with
                # heading/list_item styles, never replaces them.
                align_val = block.get("alignment")
                if align_val in _ALIGN_MAP:
                    para.alignment = _ALIGN_MAP[align_val]

                runs = block.get("runs") or []
                for run_data in runs:
                    if not isinstance(run_data, dict):
                        continue
                    text_val = run_data.get("text", "")
                    if not text_val:
                        continue
                    run = para.add_run(text_val)
                    if run_data.get("bold"):
                        run.bold = True
                    if run_data.get("italic"):
                        run.italic = True
                    if run_data.get("underline"):
                        run.underline = True
                    if run_data.get("strike"):
                        run.font.strike = True
                    font_size = run_data.get("font_size")
                    if font_size is not None:
                        try:
                            run.font.size = Pt(int(font_size))
                        except (ValueError, TypeError):
                            pass
                    # Color: stored as lowercase 6-digit hex without `#`.
                    # RGBColor.from_string requires uppercase 6-digit; we
                    # uppercase at the boundary. Defensive try/except in
                    # case malformed hex slipped past the save normalizer.
                    color_val = run_data.get("color")
                    if isinstance(color_val, str) and color_val:
                        try:
                            run.font.color.rgb = RGBColor.from_string(
                                color_val.lstrip("#").upper()
                            )
                        except (ValueError, TypeError):
                            pass
                    # Font family: setting run.font.name in python-docx
                    # writes <w:rFonts w:ascii=...>. Word falls back to a
                    # system substitute if the name isn't installed —
                    # accepted behavior (R3).
                    font_family = run_data.get("font_family")
                    if isinstance(font_family, str) and font_family:
                        run.font.name = font_family
        else:
            for para_text in chunk_text.split("\n"):
                stripped = para_text.strip()
                if stripped:
                    doc.add_paragraph(stripped)

    # ── Footer on every section ───────────────────────────────────────
    if include_footer:
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.clear()

            # Tab stops: left at 0, right at page width - margins
            page_width = section.page_width - section.left_margin - section.right_margin
            pPr = footer_para._p.get_or_add_pPr()
            tabs_elem = pPr.makeelement(qn("w:tabs"), {})
            # Right-aligned tab at the page width
            tab_right = tabs_elem.makeelement(qn("w:tab"), {
                qn("w:val"): "right",
                qn("w:pos"): str(int(page_width)),
            })
            tabs_elem.append(tab_right)
            pPr.append(tabs_elem)

            # Left side: branding text
            brand_run = footer_para.add_run("Processed by Psitta — Reading to Listening")
            brand_run.font.size = Pt(8)
            brand_run.font.italic = True
            brand_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            # Tab to right side
            tab_run = footer_para.add_run("\t")
            tab_run.font.size = Pt(8)

            # Page number field
            page_run = footer_para.add_run()
            page_run.font.size = Pt(8)
            page_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            fld_char_begin = page_run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
            page_run._r.append(fld_char_begin)

            instr_run = footer_para.add_run()
            instr_run.font.size = Pt(8)
            instr_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            instr_text = instr_run._r.makeelement(qn("w:instrText"), {})
            instr_text.text = " PAGE "
            instr_run._r.append(instr_text)

            end_run = footer_para.add_run()
            end_run.font.size = Pt(8)
            fld_char_end = end_run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
            end_run._r.append(fld_char_end)

    # ── Serialize to bytes ────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Content-to-format builders (TXT / Markdown / EPUB) ───────────────────────
#
# All consume the same `chunks` shape the DOCX builder uses:
#   {title, text_content, formatted_content}
# where formatted_content is a list of blocks:
#   {type: 'heading'|'list_item'|'paragraph', level?, list_type?, alignment?,
#    runs: [{text, bold?, italic?, underline?, strike?, ...}]}
# Chunk-level synthetic titles ('Section N') are intentionally NOT emitted —
# real structure already lives in the formatted_content headings.


def _runs_plain(runs: list | None) -> str:
    """Concatenate run texts with no markup."""
    out: list[str] = []
    for r in runs or []:
        if isinstance(r, dict) and r.get("text"):
            out.append(r["text"])
    return "".join(out)


def _runs_md(runs: list | None) -> str:
    """Render runs to inline Markdown (bold/italic/strike; underline has no
    portable Markdown so it is dropped)."""
    out: list[str] = []
    for r in runs or []:
        if not isinstance(r, dict):
            continue
        t = r.get("text", "")
        if not t:
            continue
        if r.get("strike"):
            t = f"~~{t}~~"
        if r.get("bold") and r.get("italic"):
            t = f"***{t}***"
        elif r.get("bold"):
            t = f"**{t}**"
        elif r.get("italic"):
            t = f"*{t}*"
        out.append(t)
    return "".join(out)


def _runs_xhtml(runs: list | None) -> str:
    """Render runs to inline XHTML (for EPUB), escaping text."""
    import html as _html

    out: list[str] = []
    for r in runs or []:
        if not isinstance(r, dict):
            continue
        t = r.get("text", "")
        if not t:
            continue
        t = _html.escape(t)
        if r.get("bold"):
            t = f"<strong>{t}</strong>"
        if r.get("italic"):
            t = f"<em>{t}</em>"
        if r.get("underline"):
            t = f"<u>{t}</u>"
        if r.get("strike"):
            t = f"<s>{t}</s>"
        out.append(t)
    return "".join(out)


def _heading_level(block: dict, default: int = 2) -> int:
    try:
        return max(1, min(int(block.get("level", default)), 6))
    except (ValueError, TypeError):
        return default


def _build_txt(*, title: str, chunks: list[dict]) -> bytes:
    """Plain-text export: document title, then content paragraphs."""
    lines: list[str] = [title, ""]
    for chunk in chunks:
        fmt = chunk.get("formatted_content")
        if fmt:
            for block in fmt:
                if not isinstance(block, dict):
                    continue
                t = _runs_plain(block.get("runs"))
                if t.strip():
                    lines.append(t)
                    lines.append("")
        else:
            for para in (chunk.get("text_content") or "").split("\n"):
                if para.strip():
                    lines.append(para.strip())
                    lines.append("")
    return "\n".join(lines).encode("utf-8")


def _build_markdown(*, title: str, chunks: list[dict]) -> bytes:
    """Markdown export from formatted_content (headings, lists, emphasis)."""
    lines: list[str] = [f"# {title}", ""]
    for chunk in chunks:
        fmt = chunk.get("formatted_content")
        if fmt:
            for block in fmt:
                if not isinstance(block, dict):
                    continue
                btype = block.get("type", "paragraph")
                inline = _runs_md(block.get("runs"))
                if not inline.strip():
                    continue
                if btype == "heading":
                    lines.append("#" * _heading_level(block) + " " + inline)
                elif btype == "list_item":
                    prefix = (
                        "1. " if block.get("list_type") == "numbered" else "- "
                    )
                    lines.append(prefix + inline)
                else:
                    lines.append(inline)
                lines.append("")
        else:
            for para in (chunk.get("text_content") or "").split("\n"):
                if para.strip():
                    lines.append(para.strip())
                    lines.append("")
    return "\n".join(lines).encode("utf-8")


def _build_epub(
    *, title: str, chunks: list[dict], project_name: str | None
) -> bytes:
    """Minimal, valid EPUB2 (one content document) built from the chunks."""
    import html as _html
    import uuid as _uuid
    import zipfile

    body: list[str] = []
    list_open: str | None = None

    def close_list() -> None:
        nonlocal list_open
        if list_open:
            body.append(f"</{list_open}>")
            list_open = None

    for chunk in chunks:
        fmt = chunk.get("formatted_content")
        if fmt:
            for block in fmt:
                if not isinstance(block, dict):
                    continue
                btype = block.get("type", "paragraph")
                inline = _runs_xhtml(block.get("runs"))
                if btype == "list_item":
                    want = "ol" if block.get("list_type") == "numbered" else "ul"
                    if list_open != want:
                        close_list()
                        body.append(f"<{want}>")
                        list_open = want
                    body.append(f"<li>{inline or '&#160;'}</li>")
                    continue
                close_list()
                if btype == "heading":
                    lvl = _heading_level(block)
                    if inline.strip():
                        body.append(f"<h{lvl}>{inline}</h{lvl}>")
                elif inline.strip():
                    body.append(f"<p>{inline}</p>")
        else:
            close_list()
            for para in (chunk.get("text_content") or "").split("\n"):
                if para.strip():
                    body.append(f"<p>{_html.escape(para.strip())}</p>")
    close_list()

    esc_title = _html.escape(title)
    esc_creator = _html.escape(project_name or "Psitta")
    book_id = str(_uuid.uuid4())
    body_html = "\n".join(body)

    content_xhtml = (
        '<?xml version="1.0" encoding="utf-8"?>\n'
        '<!DOCTYPE html>\n'
        '<html xmlns="http://www.w3.org/1999/xhtml">\n'
        f"<head><title>{esc_title}</title>"
        '<meta http-equiv="Content-Type" content="text/html; charset=utf-8"/>'
        "</head>\n"
        f"<body>\n<h1>{esc_title}</h1>\n{body_html}\n</body>\n</html>\n"
    )
    container_xml = (
        '<?xml version="1.0"?>\n'
        '<container version="1.0" '
        'xmlns="urn:oasis:names:tc:opendocument:xmlns:container">\n'
        '  <rootfiles><rootfile full-path="OEBPS/content.opf" '
        'media-type="application/oebps-package+xml"/></rootfiles>\n'
        "</container>\n"
    )
    opf = (
        '<?xml version="1.0" encoding="utf-8"?>\n'
        '<package xmlns="http://www.idpf.org/2007/opf" version="2.0" '
        'unique-identifier="bookid">\n'
        '  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">\n'
        f"    <dc:title>{esc_title}</dc:title>\n"
        "    <dc:language>en</dc:language>\n"
        f"    <dc:creator>{esc_creator}</dc:creator>\n"
        f'    <dc:identifier id="bookid">urn:uuid:{book_id}</dc:identifier>\n'
        "  </metadata>\n"
        "  <manifest>\n"
        '    <item id="content" href="content.xhtml" '
        'media-type="application/xhtml+xml"/>\n'
        '    <item id="ncx" href="toc.ncx" '
        'media-type="application/x-dtbncx+xml"/>\n'
        "  </manifest>\n"
        '  <spine toc="ncx"><itemref idref="content"/></spine>\n'
        "</package>\n"
    )
    ncx = (
        '<?xml version="1.0" encoding="utf-8"?>\n'
        '<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">\n'
        f'  <head><meta name="dtb:uid" content="urn:uuid:{book_id}"/></head>\n'
        f"  <docTitle><text>{esc_title}</text></docTitle>\n"
        '  <navMap><navPoint id="np1" playOrder="1">'
        f"<navLabel><text>{esc_title}</text></navLabel>"
        '<content src="content.xhtml"/></navPoint></navMap>\n'
        "</ncx>\n"
    )

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        # mimetype MUST be first and stored uncompressed.
        z.writestr(
            "mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED
        )
        z.writestr("META-INF/container.xml", container_xml)
        z.writestr("OEBPS/content.opf", opf)
        z.writestr("OEBPS/toc.ncx", ncx)
        z.writestr("OEBPS/content.xhtml", content_xhtml)
    return buf.getvalue()


def _runs_rl(runs: list | None) -> str:
    """Render runs to ReportLab Paragraph mini-markup (<b>/<i>/<u>/<strike>),
    escaping text so & < > are safe."""
    import html as _html

    out: list[str] = []
    for r in runs or []:
        if not isinstance(r, dict):
            continue
        t = r.get("text", "")
        if not t:
            continue
        t = _html.escape(t)
        if r.get("bold"):
            t = f"<b>{t}</b>"
        if r.get("italic"):
            t = f"<i>{t}</i>"
        if r.get("underline"):
            t = f"<u>{t}</u>"
        if r.get("strike"):
            t = f"<strike>{t}</strike>"
        out.append(t)
    return "".join(out)


def _build_pdf(
    *, title: str, chunks: list[dict], project_name: str | None
) -> bytes:
    """PDF export via ReportLab Platypus from the document's content."""
    import html as _html

    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    bullet_style = ParagraphStyle(
        "PsittaBullet",
        parent=styles["Normal"],
        leftIndent=20,
        bulletIndent=6,
        alignment=TA_LEFT,
    )

    flow: list = [
        Paragraph(_html.escape(title), styles["Title"]),
        Spacer(1, 12),
    ]
    if project_name:
        flow.append(Paragraph(_html.escape(project_name), styles["Heading3"]))
        flow.append(Spacer(1, 8))

    num_counter = 0
    for chunk in chunks:
        fmt = chunk.get("formatted_content")
        if fmt:
            for block in fmt:
                if not isinstance(block, dict):
                    continue
                btype = block.get("type", "paragraph")
                inline = _runs_rl(block.get("runs"))
                if not inline.strip():
                    continue
                if btype == "heading":
                    num_counter = 0
                    lvl = min(_heading_level(block), 4)
                    flow.append(Paragraph(inline, styles[f"Heading{lvl}"]))
                elif btype == "list_item":
                    if block.get("list_type") == "numbered":
                        num_counter += 1
                        bullet = f"{num_counter}."
                    else:
                        num_counter = 0
                        bullet = "•"
                    flow.append(
                        Paragraph(inline, bullet_style, bulletText=bullet)
                    )
                else:
                    num_counter = 0
                    flow.append(Paragraph(inline, styles["Normal"]))
                flow.append(Spacer(1, 5))
        else:
            num_counter = 0
            for para in (chunk.get("text_content") or "").split("\n"):
                if para.strip():
                    flow.append(
                        Paragraph(_html.escape(para.strip()), styles["Normal"])
                    )
                    flow.append(Spacer(1, 5))

    if len(flow) <= 2:
        flow.append(Paragraph("(No content)", styles["Normal"]))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        title=title,
        topMargin=inch,
        bottomMargin=inch,
        leftMargin=inch,
        rightMargin=inch,
    )
    doc.build(flow)
    return buf.getvalue()


# Export targets → (media type, file extension).
_EXPORT_FORMATS = {
    "docx": (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    ),
    "pdf": ("application/pdf", "pdf"),
    "txt": ("text/plain; charset=utf-8", "txt"),
    "md": ("text/markdown; charset=utf-8", "md"),
    "epub": ("application/epub+zip", "epub"),
}


@router.get("/{document_id}/export")
async def export_document(
    document_id: UUID,
    request: Request,
    format: str = Query("docx"),
    include_cover: bool = Query(True, alias="cover"),
    include_footer: bool = Query(True, alias="footer"),
    db: AsyncSession = Depends(get_db_session),
    user_id: UUID = Depends(get_current_user_id),
):
    """Export a document's content in the requested format (docx/txt/md/epub)."""
    fmt = (format or "docx").lower()
    if fmt not in _EXPORT_FORMATS:
        raise HTTPException(
            status_code=422, detail=f"Unsupported export format: {format}"
        )
    # Fetch document metadata
    result = await db.execute(
        text(
            "SELECT title, project_id FROM documents "
            "WHERE id = :did AND user_id = :uid AND status != 'deleted'"
        ),
        {"did": document_id, "uid": str(user_id)},
    )
    row = result.mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")

    doc_title = row["title"]
    project_id = row["project_id"]

    # Resolve project name
    project_name = None
    if project_id:
        proj_result = await db.execute(
            text("SELECT name FROM projects WHERE id = :pid"),
            {"pid": project_id},
        )
        proj_row = proj_result.mappings().first()
        if proj_row:
            project_name = proj_row["name"]

    # Fetch chunks
    chunks_result = await db.execute(
        text(
            "SELECT sequence_index, text_content, metadata_json, formatted_content "
            "FROM document_chunks WHERE document_id = :did "
            "ORDER BY sequence_index"
        ),
        {"did": document_id},
    )
    chunks = []
    for r in chunks_result.mappings():
        meta = r["metadata_json"] if isinstance(r["metadata_json"], dict) else {}
        fmt_content = (
            r["formatted_content"]
            if isinstance(r["formatted_content"], list)
            else None
        )
        chunks.append({
            "title": meta.get("title", f"Section {r['sequence_index'] + 1}"),
            "text_content": r["text_content"] or "",
            "formatted_content": fmt_content,
        })

    if not chunks:
        raise HTTPException(status_code=404, detail="No content to export")

    media_type, ext = _EXPORT_FORMATS[fmt]
    if fmt == "docx":
        data = _build_branded_docx(
            title=doc_title,
            chunks=chunks,
            project_name=project_name,
            include_cover=include_cover,
            include_footer=include_footer,
        )
    elif fmt == "pdf":
        data = _build_pdf(
            title=doc_title, chunks=chunks, project_name=project_name
        )
    elif fmt == "txt":
        data = _build_txt(title=doc_title, chunks=chunks)
    elif fmt == "md":
        data = _build_markdown(title=doc_title, chunks=chunks)
    else:  # epub
        data = _build_epub(
            title=doc_title, chunks=chunks, project_name=project_name
        )

    filename = f"{doc_title}.{ext}"
    # Content-Disposition headers must be latin-1 encodable, but titles can hold
    # em-dashes and other non-ASCII. Emit an ASCII-safe `filename=` fallback plus
    # an RFC 5987 `filename*=UTF-8''…` that modern clients use to recover the
    # real title.
    from urllib.parse import quote

    ascii_stem = doc_title.encode("ascii", "ignore").decode("ascii").strip()
    ascii_filename = f"{ascii_stem or 'document'}.{ext}"
    content_disposition = (
        f'attachment; filename="{ascii_filename}"; '
        f"filename*=UTF-8''{quote(filename)}"
    )
    await audit_service.log_event(
        db,
        action="document.export",
        resource_type="document",
        user_id=str(user_id),
        resource_id=str(document_id),
        details={
            "filename": filename,
            "format": fmt,
            "include_cover": include_cover,
            "include_footer": include_footer,
            "chunk_count": len(chunks),
        },
        ip_address=request.client.host if request.client else None,
    )
    return StreamingResponse(
        io.BytesIO(data),
        media_type=media_type,
        headers={"Content-Disposition": content_disposition},
    )


async def _eager_synthesize_chunks(
    doc_id: UUID, chunk_ids: list[UUID], user_id: UUID
) -> None:
    """Pre-synthesize all chunks in background after upload.

    user_id threads through to synthesize_with_quota so per-user EL char
    quota is enforced on eager synthesis just like on-demand calls. A
    Pro user mid-month who exhausts EL during eager batch falls through
    to Edge for the remaining chunks.
    """
    DEFAULT_VOICE_ID = "21m00Tcm4TlvDq8ikWAM"
    try:
        from psitta.db.session import async_session_factory
        from psitta.providers.tts_router import TTSRouter
        from psitta.services.audio_cache import get_mp3, put_alignment, put_mp3
        logger.info("tts.eager_synthesis.start", doc_id=str(doc_id), chunk_count=len(chunk_ids))
        async with async_session_factory() as db:
            result = await db.execute(
                text(
                    "SELECT id, text_content FROM document_chunks "
                    "WHERE id = ANY(:ids) ORDER BY sequence_index"
                ),
                {"ids": [str(cid) for cid in chunk_ids]},
            )
            chunks = result.fetchall()
        if not chunks:
            logger.warning("tts.eager_synthesis.no_chunks", doc_id=str(doc_id))
            return
        logger.info("tts.eager_synthesis.chunks_found", doc_id=str(doc_id), count=len(chunks))
        tts = TTSRouter()
        synthesized = 0
        skipped = 0
        failed = 0
        for row in chunks:
            chunk_id = str(row[0])
            chunk_text = row[1] or ""
            if not chunk_text.strip():
                skipped += 1
                continue
            cached = await get_mp3(chunk_id, DEFAULT_VOICE_ID)
            if cached:
                skipped += 1
                continue
            try:
                async with async_session_factory() as quota_db:
                    audio_bytes, alignment, provider = (
                        await tts.synthesize_with_alignment_and_quota(
                            chunk_text, DEFAULT_VOICE_ID,
                            user_id=user_id, db=quota_db,
                        )
                    )
                await put_mp3(chunk_id, DEFAULT_VOICE_ID, audio_bytes)
                # Cache the word-timing sidecar now (with-timestamps returns it
                # in the same call at no extra cost) so the FIRST read has the
                # highlight ready instead of synthesizing alignment on demand.
                if alignment is not None and provider in {"elevenlabs", "edge"}:
                    await put_alignment(chunk_id, DEFAULT_VOICE_ID, {
                        "document_id": str(doc_id),
                        "chunk_id": chunk_id,
                        "voice_id": DEFAULT_VOICE_ID,
                        "provider": provider,
                        "alignment": alignment,
                    })
                synthesized += 1
                logger.info("tts.eager_synthesis.chunk_done", doc_id=str(doc_id), chunk_id=chunk_id, size=len(audio_bytes))
            except Exception as e:
                failed += 1
                logger.warning("tts.eager_synthesis.chunk_failed", doc_id=str(doc_id), chunk_id=chunk_id, error=str(e))
        logger.info("tts.eager_synthesis.complete", doc_id=str(doc_id), synthesized=synthesized, skipped=skipped, failed=failed)
    except Exception as e:
        logger.error("tts.eager_synthesis.fatal", doc_id=str(doc_id), error=str(e))
